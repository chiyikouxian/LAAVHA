#!/usr/bin/env python3
"""Apply review fixes to main draft in place:
1. Split mixed citations [7-8,30] -> [7-8][30], [18,29] -> [18][29]
2. Add clarification about TOPSIS/hysteresis being training-time logic
3. Add model input dimension note in Section 3.1
"""
import re
from docx import Document
from docx.shared import Pt

PATH = "/home/suwen/reproduce/物联网学报_LAAVHA小论文.docx"

doc = Document(PATH)

# --- 1. Fix citation format ---
fix_count = 0
for para in doc.paragraphs:
    for run in para.runs:
        if "[7-8,30]" in run.text:
            run.text = run.text.replace("[7-8,30]", "[7-8][30]")
            fix_count += 1
        if "[18,29]" in run.text:
            run.text = run.text.replace("[18,29]", "[18][29]")
            fix_count += 1
print(f"citation fixes: {fix_count}")

# --- 2. Add TOPSIS/hysteresis clarification ---
# Insert after the algorithm flow table (Table 1), before Section 3
# Find the paragraph that starts with "3 仿真实验" and insert before it
TOPSIS_NOTE = (
    "需要指出的是，上述TOPSIS决策流程和双重滞后机制在训练阶段用于生成监督标签，"
    "即根据历史网络状态计算各候选网络的相对贴近度并施加滞后约束后，"
    "将最优网络编号作为训练目标。部署阶段，训练好的端到端神经网络直接输出目标网络编号，"
    "隐式近似了上述多步决策过程，从而在保持决策质量的同时降低了在线计算开销。"
)

# --- 3. Add model dimension note ---
MODEL_NOTE = (
    "模型输入为15维向量（3个候选网络各5项归一化指标拼接），"
    "移动状态输入为2维（速度和高度），输出为3维softmax向量对应候选网络评分。"
)

inserted_topsis = False
inserted_model = False

for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    # Insert TOPSIS note before Section 3 heading
    if t.startswith("3 仿真实验") and not inserted_topsis:
        # We need to insert a new paragraph BEFORE this one.
        # python-docx doesn't have insert_before, so we'll find the right spot
        # and add text to the paragraph just before Section 3.
        # Actually, let's find the paragraph with "综合评分趋势" summary
        # or the one right after Table 1 content.
        pass

    # Insert model note after "Python侧加载训练好的LAAVHA模型" sentence
    if "Python侧加载训练好的LAAVHA模型" in t and not inserted_model:
        # Append to this paragraph
        for run in para.runs:
            if "网络编号中，0表示5G，1表示LTE，2表示WiFi。" in run.text:
                run.text = run.text.replace(
                    "网络编号中，0表示5G，1表示LTE，2表示WiFi。",
                    "网络编号中，0表示5G，1表示LTE，2表示WiFi。" + MODEL_NOTE
                )
                inserted_model = True
                print("model note inserted")
                break

# For TOPSIS note: insert after the summary paragraph that ends Section 2
# (the one starting with "综合评分趋势")
# Actually better: insert right before "3 仿真实验" as a bridging paragraph
# We'll add it by appending to the last paragraph of section 2
for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    if "双重滞后机制进一步减少了由短时指标波动引起的频繁切换。" in t:
        # This is the summary paragraph at end of results discussion
        # Wrong place - this is in Section 3. Let me find end of Section 2.
        pass
    if "结合贴近度阈值和时间窗口滞后机制输出最终切换决策" in t:
        # This is the last row description of Table 1 - but it's in a table
        pass

# Better approach: find the paragraph right after Table 1's caption
# Table 1 caption is "表1 LAAVHA垂直切换算法流程"
# The next section heading is "3 仿真实验与结果分析"
# Between them there should be a good insertion point.
# Let's insert after the paragraph containing the algorithm summary
# that ends section 2 (before section 3).

# Find index of "3 仿真实验与结果分析"
target_idx = None
for i, para in enumerate(doc.paragraphs):
    if para.text.strip().startswith("3 仿真实验与结果分析"):
        target_idx = i
        break

if target_idx:
    # Insert a new paragraph before section 3
    # python-docx: insert paragraph by manipulating XML
    from docx.oxml.ns import qn
    new_p = doc.paragraphs[target_idx]._element
    parent = new_p.getparent()
    idx = list(parent).index(new_p)

    # Create the new paragraph element
    from copy import deepcopy
    # Use the paragraph before as template for formatting
    template_p = doc.paragraphs[target_idx - 1]._element
    new_para_elem = deepcopy(template_p)
    # Clear its runs
    for r in new_para_elem.findall(qn('w:r')):
        new_para_elem.remove(r)
    # Add our text as a single run
    from lxml import etree
    run_elem = etree.SubElement(new_para_elem, qn('w:r'))
    # Copy run properties from a body text run
    t_elem = etree.SubElement(run_elem, qn('w:t'))
    t_elem.text = TOPSIS_NOTE
    t_elem.set(qn('xml:space'), 'preserve')

    parent.insert(idx, new_para_elem)
    inserted_topsis = True
    print("TOPSIS clarification inserted before Section 3")

doc.save(PATH)
print("saved:", PATH)
print(f"Results: citations={fix_count}, topsis_note={inserted_topsis}, model_note={inserted_model}")
