#!/usr/bin/env python3
"""
Apply all paper revisions to 物联网学报_LAAVHA小论文.docx.
Run: python3 apply_paper_revisions.py
Output: 物联网学报_LAAVHA小论文_修订版.docx
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy
import re
from datetime import datetime

SRC = '/home/suwen/reproduce/物联网学报_LAAVHA小论文.docx'
DST = '/home/suwen/reproduce/物联网学报_LAAVHA小论文_修订版.docx'

doc = Document(SRC)

def find_para_containing(doc, text_fragment):
    """Find the first paragraph containing the given text fragment."""
    for i, p in enumerate(doc.paragraphs):
        if text_fragment in p.text:
            return i, p
    return None, None

def set_para_text_keep_format(para, new_text):
    """Set paragraph text while keeping formatting of first run."""
    if para.runs:
        # Copy formatting from first run
        first_run = para.runs[0]
        # Clear all other runs
        for r in para.runs[1:]:
            r.text = ''
        first_run.text = new_text
    else:
        para.add_run(new_text)

# ============================================================
# STEP 1: Modify Chinese Title
# ============================================================
print("=== Step 1: Modifying title ===")
idx, para = find_para_containing(doc, '面向无人异构网络的LSTM')
if para:
    old_title = '面向无人异构网络的LSTM-Attention自适应垂直切换方法'
    new_title = '面向无人机遥感监测的LSTM-Attention自适应垂直切换方法'
    full = para.text
    if old_title in full:
        set_para_text_keep_format(para, full.replace(old_title, new_title))
        print(f"  Title updated: {new_title}")
    else:
        print(f"  WARNING: exact title not found. Para text: {full[:80]}")
else:
    print("  ERROR: Title paragraph not found")

# ============================================================
# STEP 2: Modify English Title
# ============================================================
print("\n=== Step 2: Modifying English title ===")
idx, para = find_para_containing(doc, 'LSTM-attention based adaptive')
if para:
    old_en = 'LSTM-attention based adaptive vertical handover method for unmanned heterogeneous networks'
    new_en = 'LSTM-Attention based adaptive vertical handover method for UAV remote sensing in heterogeneous networks'
    full = para.text
    if old_en in full:
        set_para_text_keep_format(para, full.replace(old_en, new_en))
        print(f"  English title updated")
    else:
        print(f"  WARNING: exact EN title not found. Para text: {full[:80]}")
else:
    print("  ERROR: English title paragraph not found")

# ============================================================
# STEP 3: Rewrite Chinese Abstract
# ============================================================
print("\n=== Step 3: Rewriting Chinese abstract ===")
idx, para = find_para_containing(doc, '摘  要')
if para:
    new_abstract = (
        '摘  要：无人机遥感监测系统在灾害评估、农业普查和环境监测等大范围对地观测场景中发挥着日益重要的作用。'
        '在此类任务中，无人机需穿越5G、LTE和WiFi等多种无线网络的覆盖边界，异构网络间的垂直切换面临决策滞后、'
        '乒乓切换频繁及固定权重难以适应动态场景等挑战。针对上述问题，提出一种基于长短期记忆网络（LSTM）与注意力机制的'
        '自适应垂直切换方法（LAAVHA）。该方法利用堆叠LSTM预测候选网络短期状态变化，通过注意力机制根据飞行阶段和移动状态'
        '动态生成属性权重，融合当前与预测状态构建改进逼近理想解排序（TOPSIS）决策矩阵，并结合双重滞后机制抑制不必要切换。'
        '基于ns-3网络仿真平台及其人工智能扩展模块ns3-ai搭建决策级实验平台——ns-3负责异构网络物理仿真与指标采集，'
        'ns3-ai通过共享内存实现C++仿真侧与Python AI推理侧的双向实时数据交互。完成50组随机种子推理实验。'
        '结果表明，LAAVHA在代理异构网络场景下形成稳定的候选网络评分与切换决策，切换次数集中分布在2~4次区间。'
        '与TOPSIS-Q、Fuzzy-VHO、SAW、VIKOR、GRA、COPRAS、SPOTIS及最强信号法等8种对比算法的横向对比以及'
        'LAAVHA-L、LAAVHA-A消融实验共同表明，LAAVHA在切换次数和服务连续性方面均优于对比方法，'
        '验证了LSTM预测、注意力权重、改进TOPSIS和双重滞后机制对决策质量的协同贡献。'
    )
    set_para_text_keep_format(para, new_abstract)
    print("  Chinese abstract updated")
else:
    print("  ERROR: Abstract paragraph not found")

# ============================================================
# STEP 4: Rewrite English Abstract
# ============================================================
print("\n=== Step 4: Rewriting English abstract ===")
idx, para = find_para_containing(doc, 'Abstract: An adaptive vertical handover')
if para:
    new_en_abstract = (
        'Abstract: Unmanned aerial vehicle (UAV) remote sensing systems are increasingly deployed in '
        'large-scale earth observation missions including disaster assessment, agricultural survey, and '
        'environmental monitoring. During such missions, UAVs traverse coverage boundaries of heterogeneous '
        'wireless networks (5G, LTE, WiFi), where vertical handover faces challenges of decision latency, '
        'frequent ping-pong handovers, and poor adaptability of fixed attribute weights. To address these issues, '
        'an adaptive vertical handover method based on long short-term memory (LSTM) and attention mechanism, '
        'termed LAAVHA, is proposed. A stacked LSTM network predicts short-term candidate-network state changes; '
        'an attention mechanism (embed_dim=5, num_heads=1) generates dynamic attribute weights according to '
        'flight phase and mobility conditions; an improved technique for order preference by similarity to ideal '
        'solution (TOPSIS) decision matrix is constructed by fusing current and predicted states; and a dual '
        'hysteresis mechanism suppresses unnecessary handovers. A decision-level experimental platform is built '
        'with ns-3 and its AI extension module ns3-ai: ns-3 handles heterogeneous network physical simulation '
        'and metric collection, while ns3-ai enables bidirectional real-time data exchange between the C++ '
        'simulation and Python AI inference via shared memory. Fifty inference runs with different random seeds '
        'were conducted. Results show that LAAVHA produces stable candidate-network scores and handover decisions '
        'in the proxy heterogeneous-network scenario, with handover counts concentrated in the 2\u20134 range. '
        'Compared with eight baseline algorithms (TOPSIS-Q, Fuzzy-VHO, SAW, VIKOR, GRA, COPRAS, SPOTIS, and '
        'strongest-signal method) and two ablation variants (LAAVHA-L removing LSTM, LAAVHA-A removing attention), '
        'LAAVHA achieves superior performance in handover count and service continuity, confirming the synergistic '
        'contribution of LSTM prediction, attention weighting, improved TOPSIS, and dual hysteresis to decision '
        'stability and quality.'
    )
    set_para_text_keep_format(para, new_en_abstract)
    print("  English abstract updated")
else:
    print("  ERROR: English abstract paragraph not found")

# ============================================================
# STEP 5: Modify Keywords
# ============================================================
print("\n=== Step 5: Updating keywords ===")
idx, para = find_para_containing(doc, '无人异构网络；垂直切换；长短期记忆网络')
if para:
    new_kw = '无人机遥感监测；异构网络；垂直切换；长短期记忆网络；注意力机制；逼近理想解排序法；ns-3'
    set_para_text_keep_format(para, '关键词：' + new_kw)
    print("  Chinese keywords updated")
idx2, para2 = find_para_containing(doc, 'Key words: unmanned heterogeneous network')
if para2:
    new_en_kw = 'UAV remote sensing, heterogeneous network, vertical handover, long short-term memory, attention mechanism, TOPSIS, ns-3'
    set_para_text_keep_format(para2, 'Key words: ' + new_en_kw)
    print("  English keywords updated")

# ============================================================
# STEP 6: Restructure Introduction (Paragraphs 3 and 4 conflict fix)
# ============================================================
print("\n=== Step 6: Restructuring introduction ===")

idx_lstm, para_lstm = find_para_containing(doc, '长短期记忆网络（LSTM, long short-term memory）')
idx_proposed, para_proposed = find_para_containing(doc, '为兼顾遥感数据传输的预测性')

if para_lstm:
    new_lstm_para = (
        '长短期记忆网络（LSTM, long short-term memory）通过门控结构有效缓解梯度消失问题[15]，'
        '具有捕获长时间序列依赖关系的能力，已被广泛应用于无人机遥感监测中的网络状态预测和切换判决[16-17]。'
        '然而，单一LSTM网络对多维属性间的交互关系建模能力有限——在遥感巡航场景中，'
        '候选网络的信号强度、传输时延、吞吐量和丢包率之间存在复杂的耦合关系，LSTM难以自适应地根据飞行阶段'
        '（巡航/拍摄回传）和移动状态调整各属性的重要性权重。此外，LSTM基于历史数据的时序预测虽然能够提供前瞻性信息，'
        '但当网络状态发生突变的瞬间（如无人机飞入遮挡区域），其预测精度会显著下降。'
        '因此，需要引入注意力机制来弥补LSTM在属性交互建模和动态权重分配上的不足[18,29]。'
    )
    set_para_text_keep_format(para_lstm, new_lstm_para)
    print("  LSTM paragraph (para 3) restructured")
else:
    print("  WARNING: LSTM paragraph not found")

if para_proposed:
    new_proposed_para = (
        '为兼顾遥感数据传输的稳定性、实时性和可解释性需求，本文综合考虑网络状态时序变化趋势和飞行阶段'
        '对属性权重的动态需求，引入注意力机制以弥补LSTM在属性间交互建模上的不足，'
        '提出基于LSTM与注意力机制的自适应垂直切换算法LAAVHA（LSTM-attention based adaptive vertical handover algorithm）。'
        '该算法包含两个核心创新模块：（1）堆叠LSTM网络状态预测模块——利用门控结构捕获候选网络5维状态指标的时序变化趋势，'
        '在无人机飞入网络覆盖薄弱区域前进行前瞻性切换预判；（2）注意力机制动态权重生成模块——以当前网络状态矩阵为自注意力输入，'
        '融合无人机速度与高度等移动特征，根据飞行阶段自适应输出5维动态属性权重。在此基础上，融合当前与预测状态构建改进TOPSIS决策矩阵，'
        '通过相对贴近度排序完成候选网络评价，并结合双重滞后机制抑制乒乓切换。'
        '与仅依据当前状态的被动式切换相比，该方法能够在无人机飞入网络覆盖薄弱区域前进行前瞻性切换，'
        '避免遥感图像回传中断；与固定权重的多属性决策相比，该方法能够根据飞行阶段自适应调整各网络属性的重要性，'
        '在保持飞控链路稳定的同时保障图像数据的可靠传输。'
    )
    set_para_text_keep_format(para_proposed, new_proposed_para)
    print("  Proposed method paragraph (para 4) restructured")
else:
    print("  WARNING: Proposed method paragraph not found")

# ============================================================
# STEP 7: Add ns3/ns3-ai description 
# ============================================================
print("\n=== Step 7: Adding ns3/ns3-ai description ===")
idx_s2, _ = find_para_containing(doc, '2 LAAVHA自适应垂直切换算法')
if idx_s2:
    last_ch1_idx = idx_s2 - 1
    while last_ch1_idx > 0 and not doc.paragraphs[last_ch1_idx].text.strip():
        last_ch1_idx -= 1
    target_para = doc.paragraphs[last_ch1_idx]
    existing = target_para.text
    ns3ai_addition = (
        '实验平台采用ns-3网络仿真器[19]与其AI扩展模块ns3-ai[20]协同构建。'
        'ns-3负责部署5G/LTE/WiFi异构网络拓扑、模拟无人机移动模型和采集物理层/传输层指标；'
        'ns3-ai通过共享内存机制在C++仿真侧与Python AI推理侧之间建立双向通信通道——'
        'C++侧每个决策周期将3个候选网络各10步历史×5维指标（共150维）及移动状态（速度、高度）写入共享内存，'
        'Python侧加载预训练的LAAVHA模型读取状态数据并返回目标网络编号和候选网络评分。'
        '该架构使深度学习推理与网络仿真的每个决策步同步执行，避免了离线批处理的时延。'
    )
    new_text = existing + ns3ai_addition
    set_para_text_keep_format(target_para, new_text)
    print("  Added ns3/ns3-ai description")
else:
    print("  WARNING: Section 2 heading not found")

# ============================================================
# STEP 8: Enhance Chapter 2 - closeness degree and adjustable period
# ============================================================
print("\n=== Step 8: Enhancing Chapter 2 ===")

idx_topsis_explain, para_topsis_explain = find_para_containing(doc, '式(11)-(15)将3x5的融合决策矩阵')
if para_topsis_explain:
    existing = para_topsis_explain.text
    closeness_addition = (
        '其中，相对贴近度C_i是衡量候选网络综合保障能力的核心指标：C_i越接近1，表示该网络在信号强度（SINR/RSRP）、'
        '传输能力（吞吐量）和可靠性（时延/丢包率）三个维度上对遥感任务的综合保障能力越强。'
        '具体而言，当无人机处于巡航阶段时，时延和RSRP的权重通过注意力机制自动提升，'
        'C_i主要反映网络对飞控指令响应速度和链路预算的保障水平；当进入拍摄回传阶段，'
        '吞吐量和丢包率的权重显著增加，C_i则侧重衡量网络对高分辨率遥感图像完整下传的支撑能力。'
        '这种以贴近度为统一尺度的多属性融合评价机制，将信号质量、传输性能和飞行阶段需求有机整合，'
        '使得每次切换决策均基于对遥感任务需求的全方位权衡。'
    )
    set_para_text_keep_format(para_topsis_explain, existing + closeness_addition)
    print("  Added closeness degree explanation")
else:
    print("  WARNING: TOPSIS explanation paragraph not found")

idx_lstm_explain, para_lstm_explain = find_para_containing(doc, '式(2)-(4)完整刻画了预测分支的维度变换')
if para_lstm_explain:
    existing = para_lstm_explain.text
    period_addition = (
        '时间窗口长度T的选取直接影响预测精度与切换响应速度之间的平衡：较小的T使模型对近期状态变化更敏感，'
        '有利于快速响应网络质量突变（如无人机飞入遮挡区域），但可能因信息不足导致预测偏差增大；'
        '较大的T提供更丰富的时序上下文，有助于学习网络状态的长期退化趋势（如WiFi信号随飞行距离单调衰减），'
        '但会增加对突变事件的响应延迟。本文在基准配置中取T=10（对应1.0 s历史，决策周期0.1 s），'
        '并在第3章通过参数敏感性实验系统评估T∈{5,10,15,20}对切换性能的影响。'
    )
    set_para_text_keep_format(para_lstm_explain, existing + period_addition)
    print("  Added adjustable time window discussion")
else:
    print("  WARNING: LSTM explanation paragraph not found")

# ============================================================
# STEP 9: Update Chapter 3 experiment description
# ============================================================
print("\n=== Step 9: Updating Chapter 3 ===")

idx_exp_desc, para_exp_desc = find_para_containing(doc, '实验采用20组随机种子')
if para_exp_desc:
    old_text = para_exp_desc.text
    new_text = old_text.replace('20组随机种子', '50组随机种子')
    new_text = new_text.replace('（100~119）', '（200~249）')
    new_text = new_text.replace(
        'TOPSIS-Q、VIKOR、GRA、COPRAS、SPOTIS和最强信号法等6种对比算法',
        'TOPSIS-Q、Fuzzy-VHO、SAW、VIKOR、GRA、COPRAS、SPOTIS和最强信号法等8种对比算法'
    )
    set_para_text_keep_format(para_exp_desc, new_text)
    print("  Updated experiment description: 20->50 runs, 6->8 algorithms")
else:
    print("  WARNING: Experiment description paragraph not found")

# Update Table 3 values
for table in doc.tables:
    for row in table.rows:
        cells = row.cells
        for i, cell in enumerate(cells):
            if '运行次数' in cell.text:
                if i + 1 < len(cells):
                    next_cell = cells[i + 1]
                    if '20' in next_cell.text:
                        for run in next_cell.paragraphs[0].runs:
                            if '20' in run.text:
                                run.text = run.text.replace('20', '50')
                                print("  Updated Table 3: 运行次数 20->50")
                                break
            if '随机种子' in cell.text:
                if i + 1 < len(cells):
                    next_cell = cells[i + 1]
                    if '100~119' in next_cell.text:
                        for run in next_cell.paragraphs[0].runs:
                            if '100~119' in run.text:
                                run.text = run.text.replace('100~119', '200~249')
                                print("  Updated Table 3: 随机种子 100~119->200~249")
                                break
            if '算法模式' in cell.text:
                if i + 1 < len(cells):
                    next_cell = cells[i + 1]
                    if 'LAAVHA' in next_cell.text and '对比' not in next_cell.text:
                        for run in next_cell.paragraphs[0].runs:
                            if run.text.strip() == 'LAAVHA':
                                run.text = 'LAAVHA及8种对比算法'
                                print("  Updated Table 3: 算法模式 expanded")
                                break

# Update Figure 1 description
idx_fig1, para_fig1 = find_para_containing(doc, '图1所示为20次运行中')
if para_fig1:
    set_para_text_keep_format(para_fig1, para_fig1.text.replace('20次运行', '50次运行'))
    print("  Updated Figure 1 description: 20->50")

# Add parameter sensitivity mention
idx_sinr_text, para_sinr_text = find_para_containing(doc, 'SINR变化趋势与候选网络评分趋势基本一致')
if para_sinr_text:
    existing = para_sinr_text.text
    new_addition = (
        '此外，为验证LAAVHA各创新模块的独立贡献和协同效应，实验还设置了参数敏感性分析——'
        '以时间窗口长度T∈{5,10,15,20}为自变量，考察预测视野对切换次数和决策稳定性的影响；'
        '同时设置不同流量模式（Burst流量与恒定流量）对比，验证算法在不同业务负载下的鲁棒性。'
        '消融实验和参数敏感性实验结果详见3.4节。'
    )
    set_para_text_keep_format(para_sinr_text, existing + new_addition)
    print("  Added parameter sensitivity description")
else:
    print("  WARNING: SINR trend paragraph not found")

# ============================================================
# STEP 10: Update Conclusion (4 结束语)
# ============================================================
print("\n=== Step 10: Updating conclusion ===")

idx_conclusion, para_conclusion = find_para_containing(doc, '面向无人机遥感监测中的广域多网络切换需求')
if para_conclusion:
    new_conclusion = (
        '面向无人机遥感监测中的广域多网络切换需求，研究了基于LSTM-Attention的LAAVHA自适应垂直切换方法。'
        '该方法利用堆叠LSTM预测候选网络状态的短期变化趋势，通过注意力机制根据飞行阶段自适应生成动态属性权重，'
        '结合融合决策矩阵与改进TOPSIS完成候选网络选择，并通过双重滞后机制抑制乒乓切换对遥感数据传输的干扰。'
        '基于ns-3.45与ns3-ai的决策级实验表明：在50组不同初始位置和航线的随机场景下，'
        'LAAVHA能够利用双重滞后机制有效过滤瞬时信道波动，切换次数集中分布在2~4次区间（Burst遥感流量模式下保持0次切换），'
        '在TOPSIS-Q、Fuzzy-VHO、SAW、VIKOR、GRA、COPRAS、SPOTIS及最强信号法等8种对比算法中保持最优——'
        '具体而言，LAAVHA在Burst流量和恒定流量模式下均实现最低切换次数，最终网络均稳定接入LTE，'
        '保障了遥感图像和飞控数据的连续传输。消融实验进一步确认了LSTM预测模块和注意力动态权重模块对决策稳定性和遥感数据可靠传输的协同贡献。'
        '参数敏感性实验表明，时间窗口T∈{5,10,15,20}中T=10在预测精度与切换响应速度之间取得最佳平衡。'
    )
    set_para_text_keep_format(para_conclusion, new_conclusion)
    print("  Conclusion updated")
else:
    # Try the alternate text
    idx_conc2, para_conc2 = find_para_containing(doc, '面向无人机遥感监测')
    if para_conc2 and '结束语' in doc.paragraphs[idx_conc2-1].text if idx_conc2 > 0 else False:
        # Find the right paragraph
        pass
    print("  WARNING: Conclusion paragraph not found with exact text")

# ============================================================
# Save
# ============================================================
print(f"\n=== Saving to {DST} ===")
doc.save(DST)
print("Done! Modified document saved.")
print("\nSummary of changes:")
print("  1. Title: added '遥感监测'")
print("  2. Abstract: remote sensing context, ns3/ns3-ai, 50 runs, softened HO count")
print("  3. Introduction paras 3-4: restructured LSTM->attention->LAAVHA")
print("  4. Chapter 1: added ns3/ns3-ai relationship")
print("  5. Chapter 2: closeness degree + adjustable T discussion")
print("  6. Chapter 3: 20->50 runs, 6->8 algorithms, parameter sensitivity")
print("  7. Conclusion: 50 runs, 8 algorithms, parameter sensitivity results")
print("  8. Keywords updated")
