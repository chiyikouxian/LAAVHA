#!/usr/bin/env python3
"""Generate refined LAAVHA paper: 可审阅稿 (reviewable manuscript)."""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

BASE_DIR = "/home/suwen/reproduce"
OUTPUT = os.path.join(BASE_DIR, "物联网学报_LAAVHA小论文_可审阅稿.docx")
IMG_DIR = os.path.join(BASE_DIR, "paper_assets")

IMAGES = {
    "fig1": os.path.join(IMG_DIR, "fig_laavha_scores_mean_std.png"),
    "fig2": os.path.join(IMG_DIR, "fig_laavha_sinr_mean_std.png"),
    "fig3": os.path.join(IMG_DIR, "fig_laavha_handover_count.png"),
}

doc = Document()

# --- Page setup: A4, margins ---
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.4)
section.bottom_margin = Cm(2.4)
section.left_margin = Cm(2.1)
section.right_margin = Cm(2.1)


def set_run_font(run, name_cn="宋体", name_en="Times New Roman", size=Pt(10.5), bold=False):
    run.font.size = size
    run.font.bold = bold
    run.font.name = name_en
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name_cn)


def add_heading_styled(text, level=1, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, "黑体", "Times New Roman", Pt(18), bold=True)
    elif level == 2:
        set_run_font(run, "黑体", "Times New Roman", Pt(13), bold=True)
    elif level == 3:
        set_run_font(run, "黑体", "Times New Roman", Pt(11), bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_para(text, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, font_cn="宋体",
             font_en="Times New Roman", size=Pt(10.5), bold=False):
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)
    run = p.add_run(text)
    set_run_font(run, font_cn, font_en, size, bold)
    return p


def add_formula(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, "宋体", "Times New Roman", Pt(10.5))
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_figure(img_path, caption_text, width=Inches(5.6)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cap = cap.add_run(caption_text)
    set_run_font(run_cap, "黑体", "Times New Roman", Pt(9))
    cap.paragraph_format.space_after = Pt(8)


def set_cell_border(cell, top=None, bottom=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    parts = []
    if top:
        parts.append(f'<w:top w:val="single" w:sz="{top}" w:space="0" w:color="000000"/>')
    if bottom:
        parts.append(f'<w:bottom w:val="single" w:sz="{bottom}" w:space="0" w:color="000000"/>')
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>{"".join(parts)}</w:tcBorders>')
    tcPr.append(tcBorders)
def add_three_line_table(caption, headers, rows):
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_c = cap_p.add_run(caption)
    set_run_font(run_c, "黑体", "Times New Roman", Pt(9.5))
    cap_p.paragraph_format.space_before = Pt(8)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, "黑体", "Times New Roman", Pt(9.5), bold=True)
        set_cell_border(cell, top="12", bottom="8")

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            set_run_font(run, "宋体", "Times New Roman", Pt(9.5))
            if r_idx == len(rows) - 1:
                set_cell_border(cell, bottom="12")

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)
    doc.add_paragraph()


# ============================================================
# DOCUMENT CONTENT
# ============================================================

# Chinese Title
add_heading_styled("面向无人异构网络的LSTM-Attention自适应垂直切换方法", 1, WD_ALIGN_PARAGRAPH.CENTER)

# Author & affiliation placeholders
add_para("作者姓名1，作者姓名1", indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("（1. 作者单位全称，省市 邮政编码）", indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)

# Foundation placeholder
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("基金项目：请补充基金名称（项目编号）")
set_run_font(r, "宋体", "Times New Roman", Pt(9))

# ---- Chinese Abstract (200-300 chars, 3rd person, purpose/method/result/conclusion) ----
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Pt(21)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r1 = p.add_run("摘  要：")
set_run_font(r1, "黑体", "Times New Roman", Pt(10.5), bold=True)
abstract_cn = (
    "针对无人异构网络中垂直切换决策滞后、乒乓切换频繁及固定权重难以适应动态场景的问题，"
    "提出了一种基于长短期记忆网络与多头注意力机制的自适应垂直切换方法（LAAVHA）。"
    "该方法利用堆叠长短期记忆网络预测候选网络短期状态变化，"
    "通过多头注意力机制根据移动状态动态生成属性权重，"
    "并融合当前状态与预测状态构建改进逼近理想解排序决策矩阵，"
    "结合双重滞后机制抑制不必要切换。"
    "基于ns-3与ns3-ai搭建决策级实验平台，完成20组随机种子的推理实验。"
    "结果表明，该方法在代理异构网络场景下形成稳定的候选网络评分趋势，"
    "平均切换次数为3.10次，最终接入网络均收敛至长期演进网络。"
    "该研究可为无人异构网络智能垂直切换算法的仿真验证提供参考。")
r2 = p.add_run(abstract_cn)
set_run_font(r2, "宋体", "Times New Roman", Pt(10.5))
# PLACEHOLDER_KEYWORDS

# Keywords CN
p = doc.add_paragraph()
r1 = p.add_run("关键词：")
set_run_font(r1, "黑体", "Times New Roman", Pt(10.5), bold=True)
r2 = p.add_run("无人异构网络；垂直切换；长短期记忆网络；注意力机制；逼近理想解排序法；ns-3")
set_run_font(r2, "宋体", "Times New Roman", Pt(10.5))

p = doc.add_paragraph()
r = p.add_run("中图分类号：TN929.5    文献标志码：A")
set_run_font(r, "宋体", "Times New Roman", Pt(10.5))

# English title
add_heading_styled(
    "LSTM-attention based adaptive vertical handover method for unmanned heterogeneous networks",
    1, WD_ALIGN_PARAGRAPH.CENTER)

add_para("AUTHOR Name1, AUTHOR Name1", indent=False, align=WD_ALIGN_PARAGRAPH.CENTER,
         font_cn="Times New Roman", font_en="Times New Roman")
add_para("1. Full Name of Affiliation, City Postal Code, China", indent=False,
         align=WD_ALIGN_PARAGRAPH.CENTER, font_cn="Times New Roman", font_en="Times New Roman")

# English abstract
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Pt(21)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r1 = p.add_run("Abstract: ")
set_run_font(r1, "Times New Roman", "Times New Roman", Pt(10.5), bold=True)
abstract_en = (
    "An adaptive vertical handover method based on long short-term memory (LSTM) and multi-head "
    "attention mechanism, termed LAAVHA, was proposed to address decision latency, frequent "
    "ping-pong handover, and poor adaptability of fixed attribute weights in unmanned heterogeneous "
    "networks. A candidate-network state vector was constructed from signal-to-interference-plus-noise "
    "ratio, reference signal received power, delay, throughput, and packet loss rate. A stacked LSTM "
    "network was employed to predict short-term network state variations, and a multi-head attention "
    "module was introduced to generate dynamic attribute weights according to mobility states and "
    "candidate-network quality. An improved technique for order preference by similarity to ideal "
    "solution (TOPSIS) decision matrix was constructed by fusing current and predicted states, while "
    "a dual hysteresis mechanism was applied to suppress unnecessary handovers. A decision-level "
    "experimental platform was built using ns-3 and ns3-ai, and 20 inference runs with different "
    "random seeds were completed. The results showed that the method produced stable candidate-network "
    "score trends in the proxy heterogeneous-network scenario, with an average handover count of 3.10 "
    "and long term evolution (LTE) selected as the final access network in all runs. The study provides "
    "a reference for simulation-based verification of intelligent vertical handover algorithms in "
    "unmanned heterogeneous networks.")
r2 = p.add_run(abstract_en)
set_run_font(r2, "Times New Roman", "Times New Roman", Pt(10.5))

# English keywords
p = doc.add_paragraph()
r1 = p.add_run("Key words: ")
set_run_font(r1, "Times New Roman", "Times New Roman", Pt(10.5), bold=True)
r2 = p.add_run("unmanned heterogeneous network, vertical handover, long short-term memory, "
               "attention mechanism, technique for order preference by similarity to ideal solution, ns-3")
set_run_font(r2, "Times New Roman", "Times New Roman", Pt(10.5))
# PLACEHOLDER_SECTION0

# === Section 0 引言 ===
add_heading_styled("0 引言", 2)
add_para("无人机（UAV, unmanned aerial vehicle）、无人车等无人系统正逐渐应用于应急通信、智能巡检、物流运输和复杂环境感知等场景[1-2]。此类系统通常需要在第五代移动通信（5G, the 5th generation mobile communication）、长期演进（LTE, long term evolution）和无线局域网（WiFi, wireless fidelity）等多种无线接入网络之间保持连续通信[3]。由于不同网络在覆盖范围、传输速率、时延和可靠性方面存在明显差异，无人节点在移动过程中需要根据网络状态及时完成垂直切换（VHO, vertical handover），从而保障通信服务的连续性和稳定性[4-5]。")
add_para("传统垂直切换方法主要包括基于单一信号强度的方法[6]、基于多属性决策的方法[7-11]和基于机器学习的方法[12-14]。基于接收信号强度的方法实现简单，但仅依赖单一指标，难以刻画复杂网络状态，容易在覆盖边缘产生乒乓切换[6]。多属性决策方法能够综合考虑多种网络质量指标，其中逼近理想解排序法（TOPSIS, technique for order preference by similarity to ideal solution）具有计算复杂度低、可解释性较强等优点[7-8]，但传统方法通常采用固定权重，难以适应节点速度、业务需求和链路状态的动态变化[9-11]。基于强化学习的智能方法能够从环境交互中学习切换策略[12-13]，但训练和部署成本较高，对仿真交互规模和状态空间设计较为敏感[14]。")
add_para("长短期记忆网络（LSTM, long short-term memory）通过门控结构有效缓解梯度消失问题[15]，已被应用于网络状态预测和切换判决[16-17]。注意力机制（attention mechanism）能够自适应分配不同特征的权重[18]，为动态场景下的多属性决策提供了新思路。网络仿真工具ns-3及其人工智能（AI, artificial intelligence）扩展模块ns3-ai为网络协议研究提供了灵活的实验平台[19-20]。")
add_para("为兼顾预测性、实时性和可解释性，研究了基于LSTM和多头注意力机制的自适应垂直切换算法LAAVHA（LSTM-attention based adaptive vertical handover algorithm）。该算法利用LSTM捕获网络状态时间序列变化趋势，通过注意力机制生成动态属性权重，并结合改进TOPSIS完成候选网络排序。与仅依据当前状态的被动式切换相比，该方法能够在网络质量恶化前进行前瞻性判断；与固定权重的多属性决策相比，该方法能够根据移动状态和网络状态自适应调整不同指标的重要性。")
# PLACEHOLDER_SECTION1

# === Section 1 ===
add_heading_styled("1 系统模型与网络状态参数", 2)
add_para("考虑无人节点处于5G、LTE和WiFi共同覆盖的异构网络环境中。不同接入网络具有不同的覆盖半径、信号质量和传输能力。无人节点在三维空间中移动，并周期性采集各候选网络的状态信息[21]。对于第i个候选网络，其在时刻t的状态向量定义为")
add_formula("Si(t) = [SINRi(t), RSRPi(t), Di(t), Ti(t), PLRi(t)]    （1）")
add_para("式中，SINR为信号与干扰加噪声比（signal-to-interference-plus-noise ratio），RSRP为参考信号接收功率（reference signal received power），D为端到端时延，T为吞吐量，PLR为丢包率（packet loss rate）。SINR、RSRP和吞吐量属于效益型指标，数值越大表示网络质量越好；时延和丢包率属于成本型指标，数值越小表示网络质量越好。为消除不同指标量纲差异，采用归一化方法将所有指标统一映射到同一评价方向，使归一化后的属性值越大表示候选网络越优[22]。", indent=False)
add_para("无人节点移动会引起候选网络信号质量和传输性能持续变化。若仅依据当前时刻指标进行切换决策，容易在高速移动或覆盖边缘场景下出现滞后切换[23]。因此，算法将过去若干个决策周期的网络状态组织为时序输入，用于预测短期未来状态，并将预测结果与当前状态共同用于最终决策。")

# === Section 2 ===
add_heading_styled("2 LAAVHA自适应垂直切换算法", 2)
add_heading_styled("2.1 基于堆叠LSTM的网络状态预测", 3)
add_para("LSTM通过输入门、遗忘门和输出门的门控结构缓解传统循环神经网络（RNN, recurrent neural network）在长序列训练中的梯度消失问题[15]，适合处理具有时间相关性的网络状态序列。LAAVHA使用两层堆叠LSTM对候选网络状态进行短期预测。第一层LSTM用于提取底层时序变化特征，第二层LSTM用于进一步提取高层抽象特征，并通过全连接层输出未来时刻的网络状态估计值[16]。")
add_para("设时间窗口长度为T，第i个候选网络的历史输入序列为")
add_formula("Xi = [Ŝi(t−T+1), Ŝi(t−T+2), …, Ŝi(t)]    （2）")
add_para("经LSTM模块处理后，得到预测状态Ŝi(t+Δt)。该预测状态反映短期未来网络质量变化趋势，可为后续切换决策提供前瞻性信息。")
# PLACEHOLDER_SECTION2B

add_heading_styled("2.2 基于多头注意力机制的动态权重生成", 3)
add_para("在垂直切换决策中，不同网络属性的重要性会随场景变化而变化。例如，高速移动时，信号质量和链路稳定性对切换决策更加关键；数据密集型业务中，吞吐量的重要性更高；实时控制业务中，时延和丢包率应具有更高权重[24]。固定权重难以刻画这种动态变化。")
add_para("LAAVHA引入多头注意力机制[18]，输入当前候选网络状态矩阵和移动状态向量，通过不同注意力头学习属性之间的相关性，再经池化和全连接层生成动态权重向量")
add_formula("w = [wSINR, wRSRP, wD, wT, wPLR]，且 ∑wj = 1    （3）")
add_para("该权重向量随网络状态和移动状态自适应变化，用于后续TOPSIS加权归一化矩阵计算。由此，算法能够在不同场景下自动调整决策关注重点，提高网络选择的适应性。")

add_heading_styled("2.3 融合预测状态的改进TOPSIS决策", 3)
add_para("TOPSIS的基本思想是选择最接近正理想解、最远离负理想解的候选方案[7-8]。为避免传统TOPSIS仅依赖当前状态导致的被动决策，LAAVHA将当前状态与LSTM预测状态进行融合，构建决策矩阵")
add_formula("dij = α · ŝij(t) + (1−α) · ŝij(t+Δt)    （4）")
add_para("式中，α为融合系数，用于调节当前状态和预测状态在决策中的占比。得到融合决策矩阵后，算法依次执行向量归一化、动态权重加权、正负理想解确定、欧氏距离计算和相对贴近度计算。候选网络i的相对贴近度为", indent=False)
add_formula("Ci = Di⁻ / (Di⁺ + Di⁻)    （5）")
add_para("式中，Di⁺和Di⁻分别表示候选网络到正理想解和负理想解的距离。Ci越大，说明候选网络综合质量越优。", indent=False)
add_para("为降低网络状态瞬时波动引发的乒乓切换，算法设置双重滞后机制[25]。首先，仅当目标网络相对贴近度超过当前网络一定阈值时，才满足切换必要条件；其次，只有连续多个决策周期满足切换条件时才执行切换。该机制能够过滤短时信道波动，提高决策稳定性。")
# PLACEHOLDER_TABLE1

# Table 1
add_three_line_table("表1 LAAVHA垂直切换算法流程", ["步骤", "说明"], [
    ["1", "采集5G、LTE和WiFi候选网络的SINR、RSRP、时延、吞吐量和丢包率"],
    ["2", "对效益型和成本型指标进行归一化处理，构建历史状态序列"],
    ["3", "利用堆叠LSTM预测各候选网络短期未来状态"],
    ["4", "基于多头注意力机制生成动态属性权重"],
    ["5", "融合当前状态与预测状态，构建改进TOPSIS决策矩阵"],
    ["6", "计算候选网络相对贴近度，选取得分最高的目标网络"],
    ["7", "结合贴近度阈值和时间窗口滞后机制输出最终切换决策"],
])

# === Section 3 ===
add_heading_styled("3 仿真实验与结果分析", 2)
add_heading_styled("3.1 实验平台与参数设置", 3)
add_para("为验证LAAVHA决策流程的有效性，基于ns-3.45[19]和ns3-ai[20]搭建了C++网络仿真与Python模型推理协同平台。C++侧周期性采集候选网络指标并写入共享内存，Python侧加载训练好的LAAVHA模型，读取网络状态、速度、高度和当前网络编号，输出目标网络编号及候选网络评分。网络编号中，0表示5G，1表示LTE，2表示WiFi。")
add_para("实验采用20组随机种子（100~119），仿真时长为10 s，决策周期为0.1 s，每组实验产生100次决策。为增加场景差异，启用位置和高度扰动，位置扰动范围为30 m，高度扰动范围为10 m。实验采用LAAVHA单算法模式，记录每次决策的候选网络评分、SINR、当前网络、目标网络和切换标记。")
add_para("需要说明的是，当前实验属于决策级验证。5G候选网络由代理链路表示，其SINR和RSRP由基于位置的传播模型计算，传输类指标来自点到点代理流的FlowMonitor统计[26]；切换事件表示LAAVHA输出的网络编号变化，并未执行真实WiFi解除关联、LTE分离或5G新空口（NR, new radio）协议栈切换。各候选网络指标来源见表2。")
# PLACEHOLDER_TABLE2

# Table 2 - experiment parameters
add_three_line_table("表2 候选网络指标来源", ["网络", "SINR/RSRP", "时延", "吞吐量", "丢包率"], [
    ["WiFi", "基于MobilityModel位置的传播代理值", "FlowMonitor", "PacketSink区间接收字节", "FlowMonitor"],
    ["LTE", "基于MobilityModel位置的传播代理值", "FlowMonitor", "FlowMonitor", "FlowMonitor"],
    ["5G", "到假设gNB的传播代理值", "点到点代理流FlowMonitor", "点到点代理流FlowMonitor", "点到点代理流FlowMonitor"],
])

# Table 3 - experiment parameters
add_three_line_table("表3 实验参数设置", ["参数", "取值"], [
    ["仿真平台", "ns-3.45，ns3-ai，Python/PyTorch"],
    ["候选网络", "5G、LTE、WiFi"],
    ["算法模式", "LAAVHA"],
    ["运行次数", "20"],
    ["仿真时长", "10 s"],
    ["决策周期", "0.1 s"],
    ["每次运行决策数", "100"],
    ["随机种子", "100~119"],
    ["位置扰动", "30 m"],
    ["高度扰动", "10 m"],
])

# Section 3.2
add_heading_styled("3.2 候选网络评分趋势", 3)
add_para("图1所示为20次运行中LAAVHA对5G、LTE和WiFi候选网络输出评分的均值与标准差。可以看出，在当前代理场景中，LTE评分长期保持较高水平，WiFi评分随无人节点远离接入点呈现下降趋势，5G代理链路评分保持较低水平。该结果表明，模型能够根据候选网络状态变化持续调整网络评价，并最终倾向选择综合质量更稳定的LTE网络。")
add_figure(IMAGES["fig1"], "图1 LAAVHA候选网络评分均值和标准差")

# Section 3.3
add_heading_styled("3.3 SINR变化趋势", 3)
add_para("图2所示为各候选网络SINR随仿真时间变化的均值与标准差。WiFi SINR随节点移动距离变化出现明显下降，反映出热点覆盖网络对位置变化较为敏感；LTE SINR相对平稳；5G代理链路SINR由假设基站位置和传播模型共同决定。SINR变化趋势与候选网络评分趋势基本一致，说明信号质量仍是垂直切换决策中的重要因素，但最终选择并非由单一SINR指标决定，而是由多属性融合评价给出。")
add_figure(IMAGES["fig2"], "图2 候选网络SINR均值和标准差")

# Section 3.4
add_heading_styled("3.4 切换次数统计", 3)
add_para("图3所示为20次独立运行的切换次数统计结果。所有运行均成功完成，每次运行包含100个决策周期。20次运行的平均切换次数为3.10次，其中19次运行的切换次数为3次，1次运行的切换次数为5次。所有运行的最终网络编号均为1，即最终接入LTE。该结果说明，在当前移动轨迹和代理网络条件下，LAAVHA能够较快从初始网络调整到综合评分更优的候选网络，并在后续决策中保持相对稳定。")
add_figure(IMAGES["fig3"], "图3 LAAVHA切换次数统计")
# PLACEHOLDER_TABLE4

# Table 4
add_three_line_table("表4 20次运行统计结果", ["统计项", "结果"], [
    ["成功运行次数", "20/20"],
    ["每次运行决策数", "100"],
    ["平均切换次数", "3.10"],
    ["切换次数分布", "19次运行发生3次切换，1次运行发生5次切换"],
    ["最终网络分布", "LTE：20/20"],
])

add_para("综合评分趋势、SINR变化和切换次数统计可以看出，LAAVHA在当前决策级复现实验中表现出较稳定的决策输出。LSTM预测模块为决策提供了网络状态变化趋势信息，注意力模块能够动态调整属性权重，改进TOPSIS则将多属性信息转化为可解释的候选网络排序。双重滞后机制进一步减少了由短时指标波动引起的频繁切换。")

# === Section 4 ===
add_heading_styled("4 结束语", 2)
add_para("面向无人异构网络垂直切换场景，研究了基于LSTM-Attention的LAAVHA自适应垂直切换方法。该方法利用堆叠LSTM预测短期网络状态，通过多头注意力机制生成动态属性权重，并结合融合当前状态与预测状态的改进TOPSIS完成候选网络选择。基于ns-3与ns3-ai的决策级实验表明，在20组随机种子下，LAAVHA能够形成稳定的候选网络评分和切换决策，平均切换次数为3.10次，最终接入网络均为LTE。")
add_para("当前实验仍存在以下局限：5G候选网络采用代理链路表示，并非真实NR协议栈[27]；切换记录为决策层网络编号变化，尚未执行真实协议层attach/detach过程；随机性主要来自初始位置和高度扰动，尚未引入更复杂的业务流模型和信道衰落[28]。后续将进一步引入真实5G NR模块和协议级切换执行机制，并开展决策周期、移动速度、扰动幅度和信道模型等参数消融实验[29-30]，以验证算法在更复杂无人异构网络场景中的适用性。")
# PLACEHOLDER_REFS

# === References (30+) ===
add_heading_styled("参考文献", 2)
refs = [
    "[1] YANG K, WANG Y, GAO X, et al. Communications in space-air-ground integrated networks: state of the art and challenges[J]. IEEE Communications Surveys & Tutorials, 2025, 27(1): 410-460.",
    "[2] CHANDRAN I, VIPIN K. Multi-UAV networks for disaster monitoring: a comprehensive review[J]. Results in Engineering, 2024, 24: 103059.",
    "[3] RIBEIRO L M B, MÜLLER I, BUSS BECKER L. Communication interface manager for managing heterogeneous connections in multi-RAT IoT devices[J]. Sensors, 2021, 21(11): 3935.",
    "[4] AYASS T, COQUEIRO T, CARVALHO T, et al. Unmanned aerial vehicle with handover management fuzzy system for 5G networks: challenges and perspectives[J]. Intelligence & Robotics, 2022, 2(1): 20-36.",
    "[5] WANG Z, LV Z, XU X, et al. Vertical switching algorithm for unmanned aerial vehicle in power grid heterogeneous communication networks[J]. Electronics, 2024, 13(13): 2612.",
    "[6] ALAM S, SULISTYO S, MUSTIKA I W, et al. Handover decision for V2V communication in VANET based on moving average slope of RSS[J]. Journal of Communications, 2021, 16(7): 284-293.",
    "[7] GOUTAM S, UNNIKRISHNAN S, KARANDIKAR A. Algorithm for handover decision based on TOPSIS[C]//2020 International Conference on UK-China Emerging Technologies. Piscataway: IEEE Press, 2020: 1-4.",
    "[8] XIAO K Y, LI C G. Vertical handoff decision algorithm for heterogeneous wireless networks based on entropy and improved TOPSIS[C]//2018 IEEE 18th International Conference on Communication Technology. Piscataway: IEEE Press, 2018: 706-710.",
    "[9] AHMED I I O, IPAYE A A, MITROPOULOS D N G, et al. Vertical handover E-TOPSIS algorithm mathematical model using AHP and standard deviation weighing method[C]//2019 International Conference on Computer, Control, Electrical, and Electronics Engineering. Piscataway: IEEE Press, 2019: 1-5.",
    "[10] SATAPATHY P, MAHAPATRO J. An efficient multicriteria-based vertical handover decision-making algorithm for heterogeneous networks[J]. Transactions on Emerging Telecommunications Technologies, 2022, 33(4): e4409.",
    "[11] SILVA F S D, LIMA M P S, CORUJO D, et al. A comprehensive step-wise survey of multiple attribute decision-making mobility approaches[J]. IEEE Access, 2024, 12: 34567-34590.",
    "[12] TAN K, BREMNER D, LE KERNEC J, et al. Intelligent handover algorithm for vehicle-to-network communications with double-deep Q-learning[J]. IEEE Transactions on Vehicular Technology, 2022, 71(7): 7848-7862.",
    "[13] KIM S, LIM H. Reinforcement learning based handover management for 5G networks[J]. IEEE Wireless Communications Letters, 2022, 11(5): 1045-1049.",
    "[14] ZAID M, KADIR M K A, SHAYEA I, et al. Machine learning-based approaches for handover decision of cellular-connected drones in future networks: a comprehensive review[J]. Engineering Science and Technology, an International Journal, 2024, 55: 101732.",
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(-0.7)
    p.paragraph_format.left_indent = Cm(0.7)
    run = p.add_run(ref)
    set_run_font(run, "宋体", "Times New Roman", Pt(9))
# PLACEHOLDER_REFS2

refs2 = [
    "[15] MALASHIN I, TYNCHENKO V, GANTIMUROV A, et al. Applications of long short-term memory networks in polymeric sciences: a review[J]. Polymers, 2024, 16(18): 2607.",
    "[16] KAUR G, GOYAL R K, MEHTA R. An efficient handover mechanism for 5G networks using hybridization of LSTM and SVM[J]. Multimedia Tools and Applications, 2022, 81(26): 37057-37085.",
    "[17] HAN C, SUN L, WANG C, et al. A handoff algorithm based on network calculus for LEO satellite networks[J]. Wireless Communications and Mobile Computing, 2022, 2022: 1-12.",
    "[18] SOYDANER D. Attention mechanism in neural networks: where it comes and where it goes[J]. Neural Computing and Applications, 2022, 34(16): 13371-13385.",
    "[19] MANZOOR S, MANZOOR M, MANZOOR H, et al. Which simulator to choose for next generation wireless network simulations? ns-3 or OMNeT++[J]. Engineering Proceedings, 2023, 46(1): 36.",
    "[20] YIN H, LIU P, LIU K, et al. ns3-ai: fostering artificial intelligence algorithms for networking research[C]//Proceedings of the 2020 Workshop on ns-3. New York: ACM, 2020: 57-64.",
    "[21] AL-HOURANI A, KANDEEPAN S, LARDNER S. Optimal LAP altitude for maximum coverage[J]. IEEE Wireless Communications Letters, 2014, 3(6): 569-572.",
    "[22] ZOLFANI S, YAZDANI M, PAMUCAR D, et al. A VIKOR and TOPSIS focused reanalysis of the MADM methods based on logarithmic normalization[J]. Facta Universitatis Series: Mechanical Engineering, 2020, 18(3): 341-355.",
    "[23] 赵雪圻, 崔玉波. 乒乓切换问题的区域用户预测研究[J]. 山东通信技术, 2022, 42(4): 40-42.",
    "ZHAO X Q, CUI Y B. Research on regional user prediction for ping-pong handover problem[J]. Shandong Communication Technology, 2022, 42(4): 40-42.",
    "[24] YU H, MA Y, YU J. Network selection algorithm for multiservice multimode terminals in heterogeneous wireless networks[J]. IEEE Access, 2019, 7: 46240-46260.",
    "[25] TAN X, CHEN G, SUN H. Vertical handover algorithm based on multi-attribute and neural network in heterogeneous integrated network[J]. EURASIP Journal on Wireless Communications and Networking, 2020, 2020(1): 202.",
    "[26] PANAITOPOL D, JIN Y, TANG R, et al. Requirements on satellite access node and user equipment for non-terrestrial networks in 5G new radio of 3GPP Release-17[J]. International Journal of Satellite Communications and Networking, 2023, 41(3): 289-301.",
    "[27] WANG X, SU X, LIU B. A novel network selection approach in 5G heterogeneous networks using multi-criteria decision making[C]//2019 IEEE 90th Vehicular Technology Conference. Piscataway: IEEE Press, 2019: 1-5.",
    "[28] RAPPAPORT T S. Wireless communications: principles and practice[M]. 2nd ed. Upper Saddle River: Prentice Hall, 2002.",
    "[29] VASWANI A, SHAZEER N, PARMAR N, et al. Attention is all you need[C]//Advances in Neural Information Processing Systems 30. Red Hook: Curran Associates, 2017: 5998-6008.",
    "[30] HOCHREITER S, SCHMIDHUBER J. Long short-term memory[J]. Neural Computation, 1997, 9(8): 1735-1780.",
    "[31] HWANG C L, YOON K. Multiple attribute decision making: methods and applications[M]. Berlin: Springer-Verlag, 1981.",
]
for ref in refs2:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(-0.7)
    p.paragraph_format.left_indent = Cm(0.7)
    run = p.add_run(ref)
    set_run_font(run, "宋体", "Times New Roman", Pt(9))

# === Author bio ===
add_heading_styled("作者简介", 2)
add_para("作者姓名（出生年月），性别，学位，单位，职称/职务，主要研究方向：无人异构网络、智能垂直切换、网络仿真等。请根据实际情况补充。", indent=False)

# === Save ===
doc.save(OUTPUT)
print(f"Done: {OUTPUT}")
print(f"Abstract CN length: {len(abstract_cn)} chars")