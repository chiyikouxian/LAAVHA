#!/usr/bin/env python3
"""按模板体例重建源程序 docx：A4 纵向、Consolas 9pt 裸代码、无行号。

体例参照 软著模板/面向应急救援的无人机异构双网协同规划软件-源程序.docx：
  - A4 纵向，上0.96/下1.00/左右1.25 in 边距
  - 代码 Consolas 9pt，段前段后 0，单倍行距
  - 每个文件前一行 [FILE] 相对路径
  - 无标题页、无行号、无 SHA-256、无行数统计、无结束标记、
    无页眉页脚、无分页符

不改 build_docs.py，保持既有产物可复现。
必须用 /usr/bin/python3 运行（python-docx 只装在系统解释器）。
"""
import json
import pathlib

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

PROJECT = pathlib.Path("/home/suwen/reproduce")
SOURCE_DIR = PROJECT / "softcopyright"
OUTPUT_DIR = pathlib.Path(
    "/home/suwen/IBN5100/无人机自组网/软著/LAAVHA软件著作权材料")
SOFTWARE_PREFIX = "无人机遥感异构网络垂直切换智能决策软件 V1.0"

CODE_FONT = "Consolas"
CODE_PT = 9.0


def configure(doc):
    s = doc.sections[0]
    s.orientation = WD_ORIENT.PORTRAIT
    s.page_width = Inches(8.27)
    s.page_height = Inches(11.69)
    s.top_margin = Inches(0.96)
    s.bottom_margin = Inches(1.00)
    s.left_margin = Inches(1.25)
    s.right_margin = Inches(1.25)
    return doc


def add_line(doc, text, code=True):
    """加一行。code=True 用 Consolas 9pt，False 用默认字体（文件分隔行）。"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    pf.first_line_indent = 0
    if not text:
        return p
    r = p.add_run(text)
    if code:
        r.font.name = CODE_FONT
        r.font.size = Pt(CODE_PT)
        rpr = r._element.get_or_add_rPr()
        rf = rpr.get_or_add_rFonts()
        rf.set(qn("w:ascii"), CODE_FONT)
        rf.set(qn("w:hAnsi"), CODE_FONT)
        rf.set(qn("w:eastAsia"), CODE_FONT)
    return p


def refresh_inventory():
    """按磁盘现状刷新 source_inventory.json 的 lines/bytes/sha256。"""
    import hashlib

    path = SOURCE_DIR / "source_inventory.json"
    data = json.loads(path.read_text(encoding="utf-8"))
    changed = []
    for rec in data["files"]:
        f = PROJECT / rec["file"]
        raw = f.read_bytes()
        new = {
            "lines": raw.count(b"\n"),
            "bytes": len(raw),
            "sha256": hashlib.sha256(raw).hexdigest(),
        }
        if any(rec.get(k) != v for k, v in new.items()):
            changed.append((rec["file"], rec.get("lines"), new["lines"]))
            rec.update(new)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2),
                    encoding="utf-8")
    return data["files"], changed


def main():
    records, changed = refresh_inventory()
    for name, old, new in changed:
        print("清单同步 %-52s %s -> %s" % (name, old, new))

    doc = configure(Document())
    total = 0
    for rec in records:
        rel = rec["file"]
        lines = (PROJECT / rel).read_text(
            encoding="utf-8", errors="replace").splitlines()
        add_line(doc, "[FILE] %s" % rel, code=False)
        for line in lines:
            add_line(doc, line)
        add_line(doc, "")
        total += rec["lines"]

    out = OUTPUT_DIR / ("%s-源程序.docx" % SOFTWARE_PREFIX)
    doc.save(out)
    print("文件数 %d，登记行数合计 %d" % (len(records), total))
    print("saved:", out)


if __name__ == "__main__":
    main()
