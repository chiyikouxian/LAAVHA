from __future__ import annotations

import hashlib
import json
import re
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT = Path(__file__).resolve().parents[2]
SOURCE_DIR = PROJECT / "softcopyright"
OUTPUT_DIR = Path("/home/suwen/IBN5100/无人机自组网/软著/LAAVHA软件著作权材料")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
SOFTWARE_PREFIX = "无人机遥感异构网络垂直切换智能决策软件 V1.0"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Noto Sans CJK SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure_doc(doc, landscape=False):
    section = doc.sections[0]
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Noto Sans CJK SC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
    normal.font.size = Pt(10.5)
    for name, size, color in [("Title", 20, "16384D"), ("Heading 1", 15, "16384D"), ("Heading 2", 12.5, "24556B"), ("Heading 3", 11, "365A70")]:
        style = styles[name]
        style.font.name = "Noto Sans CJK SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
    if "CodeBlock" not in styles:
        style = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "DejaVu Sans Mono"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "DejaVu Sans Mono")
        style.font.size = Pt(7.2)
    return doc


def add_title(doc, title, subtitle=None):
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title)
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(subtitle)
        r.italic = True
        r.font.color.rgb = RGBColor(90, 90, 90)


def add_table(doc, rows, widths=None):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            set_cell_text(table.cell(i, j), value, bold=(i == 0))
            if i == 0:
                set_cell_shading(table.cell(i, j), "DCEAF2")
            if widths and j < len(widths):
                table.cell(i, j).width = Inches(widths[j])
    doc.add_paragraph()


def add_markdown(doc, path, skip_leading_title=False):
    lines = path.read_text(encoding="utf-8").splitlines()
    paragraph = []
    table_rows = []
    skipped_title = False

    def flush_paragraph():
        nonlocal paragraph
        if paragraph:
            text = " ".join(x.strip() for x in paragraph).strip()
            if text:
                p = doc.add_paragraph(text)
                p.paragraph_format.space_after = Pt(4)
        paragraph = []

    def flush_table():
        nonlocal table_rows
        if table_rows:
            add_table(doc, table_rows)
        table_rows = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            flush_paragraph(); flush_table(); continue
        # Markdown comments and block quotes in the editable sources are
        # drafting instructions, not part of the formal registration text.
        if stripped.startswith("<!--") or stripped.startswith(">"):
            flush_paragraph(); flush_table()
            continue
        image = re.match(r"!\[([^]]*)\]\(([^)]+)\)", stripped)
        if image:
            flush_paragraph(); flush_table()
            img = (path.parent / image.group(2)).resolve()
            if img.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(img), width=Inches(6.7))
                cap = doc.add_paragraph(image.group(1))
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].italic = True
            continue
        if stripped.startswith("#"):
            flush_paragraph(); flush_table()
            level = min(3, len(stripped) - len(stripped.lstrip("#")))
            text = stripped[level:].strip()
            if skip_leading_title and not skipped_title and level == 1:
                skipped_title = True
                continue
            try:
                doc.add_heading(text, level=level)
            except KeyError:
                # Some localized DOCX templates do not expose the English
                # built-in heading names used by python-docx.
                doc.add_paragraph(text)
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            flush_paragraph()
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if all(re.fullmatch(r"[-: ]+", c or "-") for c in cells):
                continue
            table_rows.append(cells)
            continue
        if stripped.startswith("- [") or stripped.startswith("- "):
            flush_paragraph(); flush_table()
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(re.sub(r"^- (\[[ xX]\] )?", "", stripped))
            continue
        paragraph.append(stripped)
    flush_paragraph(); flush_table()


def build_summary():
    doc = configure_doc(Document())
    add_title(doc, "计算机软件著作权登记内容摘要", SOFTWARE_PREFIX)
    add_markdown(doc, SOURCE_DIR / "summary.md", skip_leading_title=True)
    doc.save(OUTPUT_DIR / f"{SOFTWARE_PREFIX}-内容摘要.docx")


def build_application():
    doc = configure_doc(Document())
    add_title(doc, "计算机软件著作权登记申请表（内容草案）")
    add_markdown(doc, SOURCE_DIR / "application_form.md", skip_leading_title=True)
    doc.save(OUTPUT_DIR / f"{SOFTWARE_PREFIX}-申请表.docx")


def build_design():
    doc = configure_doc(Document())
    add_title(doc, SOFTWARE_PREFIX, "软件设计说明书 V1.0（草案）")
    add_markdown(doc, SOURCE_DIR / "design_description.md", skip_leading_title=True)
    doc.save(OUTPUT_DIR / f"{SOFTWARE_PREFIX}-设计说明书v1.0.docx")


# Fixed registration scope: all current authored functional code.  Material
# generation scripts, generated data, model weights, caches and deprecated
# files are deliberately excluded.
CORE_FILES = [
    "CMakeLists_laavha.txt",
    "LAAVHA改进算法训练程序.py",
    "fuzzy_vho.py",
    "laavha-handover.cc",
    "laavha_batch_runner.py",
    "laavha_inference.py",
    "laavha_msg.h",
    "laavha_plot.py",
    "laavha_py.cc",
    "madm_comparison.py",
    "make_pub_figures.py",
    "regenerate_figures.py",
    "saw_madm.py",
    "topsis_q.py",
    "softcopyright/tools/laavha_viz/__init__.py",
    "softcopyright/tools/laavha_viz/trace_model.py",
    "softcopyright/tools/laavha_viz/surface.py",
    "softcopyright/tools/laavha_viz/render.py",
    "softcopyright/tools/laavha_viz/app.py",
    "softcopyright/tools/laavha_viz/__main__.py",
    "experiments/enhanced_proof_experiments.py",
    "experiments/exp_a_adaptive_hysteresis.py",
    "experiments/gen_fig5_6.py",
    "experiments/generate_nature_figures.py",
    "experiments/generate_network_coverage_en.py",
    "experiments/parameter_sensitivity.py",
    "experiments/stress_5g_degradation.py",
]


def source_inventory():
    records = []
    for rel in CORE_FILES:
        path = PROJECT / rel
        data = path.read_bytes()
        records.append({
            "file": rel,
            # Match the conventional `wc -l` source-program count used in
            # the application form, including its treatment of a final line
            # without a terminating newline.
            "lines": data.count(b"\n"),
            "bytes": len(data),
            "sha256": hashlib.sha256(data).hexdigest(),
        })
    out = SOURCE_DIR / "source_inventory.json"
    out.write_text(json.dumps({"project": str(PROJECT), "files": records}, ensure_ascii=False, indent=2), encoding="utf-8")
    return records


def add_code_block(doc, rel, start=None, end=None):
    path = PROJECT / rel
    lines = path.read_text(encoding="utf-8", errors="replace").splitlines()
    if start is None:
        start = 1
    if end is None:
        end = len(lines)
    p = doc.add_paragraph()
    r = p.add_run(f"文件：{rel}（第{start}—{end}行）")
    r.bold = True
    r.font.name = "Noto Sans CJK SC"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
    for number, line in enumerate(lines[start - 1:end], start=start):
        p = doc.add_paragraph(style="CodeBlock")
        p.paragraph_format.space_after = Pt(0)
        p.add_run(f"{number:04d} | {line}")


def build_source_listing():
    records = source_inventory()
    doc = configure_doc(Document(), landscape=True)
    total_lines = sum(record["lines"] for record in records)
    add_title(doc, "LAAVHA软件源程序文档", "完整源程序")
    p = doc.add_paragraph(f"本文件按固定顺序收录21个现行功能代码文件，共{total_lines}行。模型权重、训练数据、NS-3外部工作区、软著材料生成脚本、编译生成文件及弃用文件不作为源代码正文。")
    p.paragraph_format.space_after = Pt(6)
    for idx, rec in enumerate(records, start=1):
        doc.add_heading(f"{idx}. {rec['file']}", level=2)
        display_lines = len((PROJECT / rec["file"]).read_text(encoding="utf-8", errors="replace").splitlines())
        line_text = f"行数（wc -l）：{rec['lines']}"
        if display_lines != rec["lines"]:
            line_text += f"；正文显示行数：{display_lines}（文件末尾无换行符）"
        doc.add_paragraph(f"{line_text}；字节数：{rec['bytes']}；SHA-256：{rec['sha256']}")
        add_code_block(doc, rec["file"])
    doc.save(OUTPUT_DIR / f"{SOFTWARE_PREFIX}-源程序.docx")


def build_fallback_listing():
    records = source_inventory()
    doc = configure_doc(Document(), landscape=True)
    add_title(doc, "LAAVHA软件源程序文档", "首30页加末30页登记备选版（草案）")
    doc.add_paragraph("本文件仅在完整源程序渲染超过60页时作为登记备选版使用。页面截取应以最终PDF页码为准；此处保留文件边界和截取说明。")
    for idx, rec in enumerate(records, start=1):
        doc.add_heading(f"{idx}. {rec['file']}", level=2)
        lines = (PROJECT / rec["file"]).read_text(encoding="utf-8", errors="replace").splitlines()
        keep = list(range(min(15, len(lines)))) + list(range(max(15, len(lines) - 15), len(lines)))
        p = doc.add_paragraph(f"文件行数：{len(lines)}；本备选源按文件头部和尾部保留，共展示{len(keep)}行。")
        for number in keep:
            p = doc.add_paragraph(style="CodeBlock")
            p.paragraph_format.space_after = Pt(0)
            p.add_run(f"{number+1:04d} | {lines[number]}")
        doc.add_page_break()
    doc.save(OUTPUT_DIR / f"{SOFTWARE_PREFIX}-源程序行级预览.docx")


def main():
    build_summary()
    build_application()
    build_design()
    build_source_listing()
    build_fallback_listing()
    shutil.copy2(SOURCE_DIR / "source_inventory.json", OUTPUT_DIR / "source_inventory.json")
    print(f"Generated documents in {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
