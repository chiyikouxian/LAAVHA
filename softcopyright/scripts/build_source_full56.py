#!/usr/bin/env python3
"""按提交源文件目录（56 个文件）重建源程序 docx，模板体例、A4 纵向。

文件顺序：登记的 27 个在前（保持原顺序），其后依次为
运行可视化界面（浏览器端）、软著材料生成脚本、文档处理与统计工具。
体例与 build_source_plain.py 一致：Consolas 9pt 裸代码、[FILE] 路径分隔行、
无行号、无标题页、无结束标记。

必须用 /usr/bin/python3 运行。
"""
import json
import pathlib

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

SUB = pathlib.Path("/home/suwen/IBN5100/无人机自组网/软著/"
                   "LAAVHA软件著作权材料/提交源文件")
OUT_DIR = SUB.parent
PREFIX = "无人机遥感异构网络垂直切换智能决策软件 V1.0"
INV = pathlib.Path("/home/suwen/reproduce/softcopyright/source_inventory.json")

CODE_FONT = "Consolas"
CODE_PT = 9.0
MANIFEST_NAME = "提交源文件清单.md"

# 登记文件在提交目录中的落地名 -> 原登记相对路径（用于排序与显示）
SUBDIR_ORDER = ["viz_web", "build_scripts", "tools"]


def configure(doc):
    s = doc.sections[0]
    s.orientation = WD_ORIENT.PORTRAIT
    s.page_width = Inches(8.27)
    s.page_height = Inches(11.69)
    s.top_margin = Inches(0.96)
    s.bottom_margin = Inches(1.00)
    s.left_margin = Inches(1.25)
    s.right_margin = Inches(1.25)
    return doc


def add_line(doc, text, code=True):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    pf.first_line_indent = 0
    if not text:
        return p
    r = p.add_run(text)
    if code:
        r.font.name = CODE_FONT
        r.font.size = Pt(CODE_PT)
        rpr = r._element.get_or_add_rPr()
        rf = rpr.get_or_add_rFonts()
        rf.set(qn("w:ascii"), CODE_FONT)
        rf.set(qn("w:hAnsi"), CODE_FONT)
        rf.set(qn("w:eastAsia"), CODE_FONT)
    return p


def ordered_files():
    """返回 [(显示路径, 磁盘路径)]，登记 27 个在前。"""
    inv = json.loads(INV.read_text(encoding="utf-8"))
    out = []
    seen = set()
    for rec in inv["files"]:
        rel = rec["file"]
        name = pathlib.Path(rel).name
        if rel.startswith("softcopyright/tools/laavha_viz/"):
            disk = SUB / "laavha_viz" / name
            show = "laavha_viz/" + name
        else:
            disk = SUB / name
            show = name
        if not disk.exists():
            raise SystemExit("提交目录缺少登记文件: %s" % disk)
        out.append((show, disk))
        seen.add(disk.resolve())

    for sub in SUBDIR_ORDER:
        base = SUB / sub
        if not base.is_dir():
            continue
        for p in sorted(base.rglob("*")):
            if not p.is_file() or p.resolve() in seen:
                continue
            out.append((str(p.relative_to(SUB)), p))
            seen.add(p.resolve())

    # 兜底：提交目录里其余未收的文件（排除清单本身）
    for p in sorted(SUB.rglob("*")):
        if not p.is_file() or p.resolve() in seen:
            continue
        if p.name == MANIFEST_NAME:
            continue
        out.append((str(p.relative_to(SUB)), p))
        seen.add(p.resolve())
    return out


def main():
    files = ordered_files()
    doc = configure(Document())
    total = 0
    for show, disk in files:
        lines = disk.read_text(encoding="utf-8", errors="replace").splitlines()
        add_line(doc, "[FILE] %s" % show, code=False)
        for line in lines:
            add_line(doc, line)
        add_line(doc, "")
        total += disk.read_bytes().count(b"\n")

    out = OUT_DIR / ("%s-源程序.docx" % PREFIX)
    doc.save(out)
    print("收录 %d 个文件，%d 行（wc -l 口径）" % (len(files), total))
    print("saved:", out)


if __name__ == "__main__":
    main()
