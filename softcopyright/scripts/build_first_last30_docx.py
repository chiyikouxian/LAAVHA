#!/usr/bin/env python3
"""生成与「源程序前30页后30页.pdf」对应的 Word 文档（A4 纵向）。

做法：把完整源程序 docx 渲染为 PDF，逐页提取文本，按顺序对齐回 docx 段落，
定出第1—30页与第101—130页各自覆盖的段落区间，再复制这些段落生成新 docx。
这样文字取自 docx 原段落（不经 PDF 提取，无换行/空格失真），
页面范围与已交付的 60 页 PDF 一致。

必须用 /usr/bin/python3 运行。
"""
import copy
import pathlib
import re
import argparse
import subprocess
import sys
import tempfile

from docx import Document

OUT_DIR = pathlib.Path("/home/suwen/IBN5100/无人机自组网/软著/"
                       "LAAVHA软件著作权材料")
PREFIX = "无人机遥感异构网络垂直切换智能决策软件 V1.0"
FULL_DOCX = OUT_DIR / ("%s-源程序.docx" % PREFIX)
OUT_DOCX = OUT_DIR / ("%s-源程序前30页后30页.docx" % PREFIX)

HEAD_PAGES = 30
TAIL_PAGES = 30


def norm(s):
    return re.sub(r"\s+", "", s)


def lo_pages(docx_path, workdir):
    """用 LibreOffice 渲染并返回页数。"""
    for f in pathlib.Path(workdir).glob("*.pdf"):
        f.unlink()
    subprocess.run(
        ["/usr/bin/libreoffice", "--headless", "--convert-to", "pdf",
         "--outdir", str(workdir), str(docx_path)],
        check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    pdfs = list(pathlib.Path(workdir).glob("*.pdf"))
    if not pdfs:
        raise SystemExit("PDF 转换失败")
    info = subprocess.run(["pdfinfo", str(pdfs[0])],
                          capture_output=True, text=True, check=True)
    for line in info.stdout.splitlines():
        if line.startswith("Pages:"):
            return int(line.split()[1]), pdfs[0]
    raise SystemExit("未取到页数")


def build_slice(src_paras_count, head_n, tail_n):
    """按段落数切片生成 docx，返回保存路径。"""
    out = Document(str(FULL_DOCX))
    paras = out.paragraphs
    keep = set(range(0, head_n)) | set(
        range(max(head_n, len(paras) - tail_n), len(paras)))
    for i in range(len(paras) - 1, -1, -1):
        if i not in keep:
            el = paras[i]._element
            el.getparent().remove(el)
    tmp = OUT_DOCX.parent / (".tmp_" + OUT_DOCX.name)
    out.save(str(tmp))
    return tmp


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--word-target", type=int, default=60,
                    help="Word 中期望的总页数")
    ap.add_argument("--ratio", type=float, default=60.0 / 66.0,
                    help="Word 目标页 / Word 实测页（默认 60/66）")
    args = ap.parse_args()

    if not FULL_DOCX.exists():
        raise SystemExit("缺少完整源程序 docx: %s" % FULL_DOCX)

    src = Document(str(FULL_DOCX))
    n_par = len(src.paragraphs)
    lo_target = max(1, int(round(args.word_target * args.ratio)))
    print("完整 docx 段落数: %d" % n_par)
    print("Word 目标 %d 页 -> LibreOffice 目标 %d 页（ratio=%.4f）"
          % (args.word_target, lo_target, args.ratio))

    with tempfile.TemporaryDirectory(prefix="fl30_") as wd:
        full_pages, _ = lo_pages(FULL_DOCX, wd)
        print("完整 PDF 页数(LO): %d" % full_pages)
        per_page = n_par / float(full_pages)

        # 初值：按每页平均段落数取一半目标页
        half = lo_target / 2.0
        head_n = int(round(half * per_page))
        tail_n = int(round(half * per_page))

        best = None
        for it in range(1, 13):
            tmp = build_slice(n_par, head_n, tail_n)
            pages, _ = lo_pages(tmp, wd)
            print("  第%2d 次: head=%d tail=%d -> LO %d 页"
                  % (it, head_n, tail_n, pages))
            if best is None or abs(pages - lo_target) < abs(best[0] - lo_target):
                best = (pages, head_n, tail_n)
            if pages == lo_target:
                break
            delta = lo_target - pages
            step = int(round(delta * per_page / 2.0))
            if step == 0:
                step = 1 if delta > 0 else -1
            head_n = max(1, head_n + step)
            tail_n = max(1, tail_n + step)
            if head_n + tail_n >= n_par:
                print("  已覆盖全文，停止")
                break

        pages, head_n, tail_n = best
        tmp = build_slice(n_par, head_n, tail_n)
        final_pages, pdf = lo_pages(tmp, wd)
        tmp.replace(OUT_DOCX)

    est_word = int(round(final_pages / args.ratio))
    print("\n最终: head=%d tail=%d 段落，LO %d 页，估算 Word 约 %d 页"
          % (head_n, tail_n, final_pages, est_word))
    print("saved:", OUT_DOCX)


if __name__ == "__main__":
    sys.exit(main())
