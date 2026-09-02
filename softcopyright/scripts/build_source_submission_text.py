from pathlib import Path
import subprocess
import tempfile
import zipfile
import re

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt
from docx.oxml.ns import qn


SOURCE = Path("/home/suwen/IBN5100/无人机自组网/软著/LAAVHA软件著作权材料/无人机遥感异构网络垂直切换智能决策软件 V1.0-源程序.docx")
OUTPUT = SOURCE.with_name(SOURCE.stem + "-前后30页文本版.docx")


def extract_pages(pdf_path: Path, first: int, last: int):
    text = subprocess.check_output(
        ["pdftotext", "-layout", "-f", str(first), "-l", str(last), str(pdf_path), "-"],
        text=True,
    )
    pages = text.split("\f")
    # pdftotext normally appends one form feed after the final page.
    if pages and not pages[-1].strip():
        pages.pop()
    expected = last - first + 1
    if len(pages) != expected:
        raise RuntimeError(f"expected {expected} pages, got {len(pages)} for {first}-{last}")
    return pages


def main():
    with tempfile.TemporaryDirectory(prefix="laavha-source-") as tmp:
        pdf = Path(tmp) / "source.pdf"
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf", "--outdir", tmp, str(SOURCE)],
            check=True,
            stdout=subprocess.DEVNULL,
        )
        generated_pdf = Path(tmp) / (SOURCE.stem + ".pdf")
        if not generated_pdf.exists():
            raise FileNotFoundError(generated_pdf)
        pages = extract_pages(generated_pdf, 1, 30) + extract_pages(generated_pdf, 117, 146)

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)

    normal = doc.styles["Normal"]
    normal.font.name = "Courier New"
    normal.font.size = Pt(7)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")

    for index, page in enumerate(pages):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(8)
        paragraph.paragraph_format.widow_control = False
        run = paragraph.add_run(page.rstrip("\n"))
        run.font.name = "Courier New"
        run.font.size = Pt(7)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
        if index < len(pages) - 1:
            paragraph.add_run().add_break(WD_BREAK.PAGE)

    doc.core_properties.title = "LAAVHA软件源程序文档（前后30页文本版）"
    doc.core_properties.subject = "软件著作权源程序提交文本"
    doc.save(OUTPUT)
    # python-docx may emit a zoom element without the required percent value.
    # Add the standard value so Office validators accept the generated package.
    with tempfile.TemporaryDirectory(prefix="laavha-docx-fix-") as tmp:
        fixed = Path(tmp) / OUTPUT.name
        with zipfile.ZipFile(OUTPUT, "r") as zin, zipfile.ZipFile(fixed, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/settings.xml":
                    data = re.sub(
                        rb"(<w:zoom\b(?![^>]*\bw:percent=)[^>]*?)\s*/>",
                        rb'\1 w:percent="100"/>',
                        data,
                        count=1,
                    )
                zout.writestr(item, data)
        fixed.replace(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
