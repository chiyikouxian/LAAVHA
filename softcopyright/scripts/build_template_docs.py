from __future__ import annotations

import datetime as dt
import re
import shutil
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from build_docs import CORE_FILES, PROJECT, add_markdown, source_inventory


TEMPLATE_DIR = Path("/home/suwen/IBN5100/无人机自组网/软著/软著模板")
OUTPUT_DIR = Path("/home/suwen/IBN5100/无人机自组网/软著/LAAVHA软件著作权材料")
# This generated design document has the preferred cover, whitespace, and
# page geometry. Only its sample body is replaced during regeneration.
DESIGN_LAYOUT_TEMPLATE = OUTPUT_DIR.parent / "LAAVHA软件著作权材料_backup_20260901_154233" / (
    "无人机遥感异构网络垂直切换智能决策软件 V1.0-设计说明书v1.0.docx"
)
SOFTWARE_NAME = "无人机遥感异构网络垂直切换智能决策软件 V1.0"
SOURCE_PROGRAM_PAGE_COUNT = 97


def patch_template_headers(path: Path, *, total_pages: int | None = None):
    """Replace template-only header text while preserving the header fields/layout."""
    replacement_total = total_pages if total_pages is not None else SOURCE_PROGRAM_PAGE_COUNT
    temp = path.with_suffix(path.suffix + ".patched")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temp, "w") as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename.startswith("word/header") and item.filename.endswith(".xml"):
                text = data.decode("utf-8")
                # The template stores the name and version in separate runs.
                text = text.replace("面向应急救援的无人机异构双网协同规划软件v", SOFTWARE_NAME[:-3])
                text = text.replace("共257页", f"共{replacement_total}页")
                data = text.encode("utf-8")
            target.writestr(item, data)
    temp.replace(path)


def clear_body(doc):
    body = doc._element.body
    sect_pr = body.find(qn("w:sectPr"))
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def clear_body_from_child(doc, first_child_index: int):
    """Keep the template front matter and remove only its sample body."""
    body = doc._element.body
    sect_pr = body.find(qn("w:sectPr"))
    for child in list(body)[first_child_index:]:
        if child is not sect_pr:
            body.remove(child)


def clear_body_after_paragraph(doc, marker: str):
    """Remove sample content after a marker paragraph, preserving front matter."""
    paragraph = next(p for p in doc.paragraphs if p.text.strip() == marker)
    body = doc._element.body
    sect_pr = body.find(qn("w:sectPr"))
    children = list(body)
    start = children.index(paragraph._p)
    for child in children[start:]:
        if child is not sect_pr:
            body.remove(child)


def set_run_font(run, name="宋体", size=10, color="FF0000", bold=False):
    run.font.name = name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def clear_paragraph_runs(paragraph):
    for child in list(paragraph._p):
        if child.tag == qn("w:r") or child.tag == qn("w:hyperlink"):
            paragraph._p.remove(child)


def set_paragraph_text(paragraph, text, *, color="000000", size=10, bold=False, font="宋体"):
    clear_paragraph_runs(paragraph)
    parts = str(text).split("\n")
    for i, part in enumerate(parts):
        if i:
            paragraph.add_run().add_break()
        run = paragraph.add_run(part)
        set_run_font(run, name=font, size=size, color=color, bold=bold)


def set_cell_text(cell, text, *, color="FF0000", size=9, font="宋体"):
    if not cell.paragraphs:
        cell.add_paragraph()
    first = cell.paragraphs[0]
    set_paragraph_text(first, text, color=color, size=size, font=font)
    for paragraph in cell.paragraphs[1:]:
        clear_paragraph_runs(paragraph)


def recursive_tables(table):
    yield table
    for row in table.rows:
        for cell in row.cells:
            for nested in cell.tables:
                yield from recursive_tables(nested)


def normalized(text):
    return " ".join(str(text).replace("\n", " ").split())


def find_summary_section(name):
    text = (PROJECT / "softcopyright/summary.md").read_text(encoding="utf-8")
    match = re.search(rf"## .*{re.escape(name)}.*?\n\n(.*?)(?=\n## |\Z)", text, re.S)
    return match.group(1).strip() if match else ""


def backup_current():
    stamp = dt.datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = OUTPUT_DIR.parent / f"{OUTPUT_DIR.name}_backup_{stamp}"
    shutil.copytree(OUTPUT_DIR, backup)
    return backup


def remove_old_generated_files():
    for path in OUTPUT_DIR.iterdir():
        if path.is_file() and path.suffix.lower() in {".docx", ".pdf"}:
            path.unlink()


def build_summary():
    target = OUTPUT_DIR / f"【填写要求】{SOFTWARE_NAME}-内容摘要.docx"
    template = TEMPLATE_DIR / "【填写要求】根据字数要求总结内容(这个文档也是需要填写的).docx"
    shutil.copy2(template, target)
    doc = Document(target)
    fields = {
        1: find_summary_section("开发目的"),
        3: find_summary_section("面向行业/领域"),
        5: find_summary_section("软件的主要功能"),
        7: find_summary_section("技术特点"),
    }
    for index, text in fields.items():
        set_paragraph_text(doc.paragraphs[index], text, color="000000", size=10)
    doc.save(target)


def fill_labeled_row(table, values):
    for row in table.rows:
        cells = row.cells
        for i, cell in enumerate(cells):
            label = normalized(cell.text)
            if label in values and i + 1 < len(cells):
                set_cell_text(cells[i + 1], values[label])


def fill_following_row(table, values):
    for ri, row in enumerate(table.rows[:-1]):
        labels = [normalized(c.text) for c in row.cells]
        if not any(label in values for label in labels):
            continue
        next_row = table.rows[ri + 1].cells
        for i, label in enumerate(labels):
            if label in values and i < len(next_row):
                set_cell_text(next_row[i], values[label])


def build_application():
    target = OUTPUT_DIR / f"{SOFTWARE_NAME}-申请表.docx"
    template = TEMPLATE_DIR / "面向应急救援的无人机异构双网协同规划软件-申请表.docx"
    shutil.copy2(template, target)
    doc = Document(target)
    main_function = find_summary_section("软件的主要功能")
    tech = find_summary_section("技术特点")
    label_values = {
        "软件名称": SOFTWARE_NAME,
        "版本号": "V1.0",
        "软件简称": "LAAVHA-VHO",
        "分类号": "[待确认]",
        "开发完成日期": "[待确认]",
        "开发方式": "● 独立开发（待确认）  ○ 合作开发  ○ 委托开发  ○ 下达任务开发",
        "姓名或名称": "[待确认]",
        "类别": "[待确认]",
        "证件类型": "[待确认]",
        "证件号码": "[待确认]",
        "国籍": "[待确认]",
        "省份/城市": "[待确认]",
        "申请方式": "●由著作权人申请  ○由代理人申请",
        "电话": "[待确认]",
        "详细地址": "[待确认]",
        "邮编": "[待确认]",
        "联系人": "[待确认]",
        "手机": "[待确认]",
        "E-mail": "[待确认]",
        "传真": "[待确认]",
    }
    for table in doc.tables:
        for nested in recursive_tables(table):
            fill_labeled_row(nested, label_values)
            fill_following_row(nested, label_values)
            for row in nested.rows:
                labels = [normalized(c.text) for c in row.cells]
                if "发表状态" in labels:
                    idx = labels.index("发表状态")
                    if idx + 1 < len(row.cells):
                        set_cell_text(row.cells[idx + 1], "○已发表（如已发表请补填日期和地点）\n●未发表（待确认）", size=8)
                if "软件环境" in labels:
                    idx = labels.index("软件环境")
                    if idx + 1 < len(row.cells):
                        set_cell_text(row.cells[idx + 1], "Ubuntu 20.04+；NS-3.45；Python 3.10+；PyTorch；NumPy/Pandas；Matplotlib；ns3-ai", size=8)
                if "硬件环境" in labels:
                    idx = labels.index("硬件环境")
                    if idx + 1 < len(row.cells):
                        set_cell_text(row.cells[idx + 1], "CPU 2 GHz以上，内存8 GB以上，硬盘20 GB以上", size=8)
                if "编程语言" in labels:
                    idx = labels.index("编程语言")
                    if idx + 1 < len(row.cells):
                        set_cell_text(row.cells[idx + 1], "Python、C++", size=9)
                if "源程序量" in labels:
                    idx = labels.index("源程序量")
                    if idx + 1 < len(row.cells):
                        total = sum(len((PROJECT / f).read_text(encoding="utf-8", errors="replace").splitlines()) for f in CORE_FILES)
                        set_cell_text(row.cells[idx + 1], f"约{total}行（核心源程序）", size=8)
                if "主要功能 和技术特点" in labels:
                    idx = labels.index("主要功能 和技术特点")
                    description = main_function + "\n\n技术特点：" + tech
                    for cell in row.cells[idx + 1:]:
                        set_cell_text(cell, description, size=7.5)
    doc.save(target)


def add_template_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run_font(run, name="黑体", size=18, color="000000", bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    set_run_font(run, name="宋体", size=12, color="000000")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("申请单位：[待确认]    [待确认]年[待确认]月[待确认]日")
    set_run_font(run, name="宋体", size=10, color="000000")


def replace_template_toc(doc, entries):
    """Replace the template TOC content while retaining its dedicated TOC page."""
    toc = next((child for child in doc._element.body if child.tag == qn("w:sdt")), None)
    if toc is None:
        return
    content = toc.find(qn("w:sdtContent"))
    if content is None:
        return
    for child in list(content):
        content.remove(child)

    title_style = "TOC 标题1" if "TOC 标题1" in [s.name for s in doc.styles] else "Normal"
    title = doc.add_paragraph("目 录", style=title_style)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    content.append(title._p)
    style_names = {style.name for style in doc.styles}
    for text, level, page in entries:
        preferred = "toc 1" if level == 1 else "toc 2"
        style = preferred if preferred in style_names else "Normal"
        paragraph = doc.add_paragraph(style=style)
        set_paragraph_text(paragraph, f"{text}{'.' * 58}{page}", color="000000", size=10)
        content.append(paragraph._p)


def apply_template_body_format(doc, start_index, body_template, section_template, subsection_template, table_template):
    """Apply the template's direct paragraph and table formatting to new content."""
    for paragraph in doc.paragraphs[start_index:]:
        if paragraph.style.name == "Heading 2":
            source = section_template
            font, size, bold = "黑体", 16, True
        elif paragraph.style.name == "Heading 3":
            source = subsection_template
            font, size, bold = "黑体", 16, False
        else:
            source = body_template
            font, size, bold = "宋体", 10, False
        paragraph._p.get_or_add_pPr().clear()
        paragraph._p.insert(0, deepcopy(source._p.pPr))
        # Markdown already supplies heading numbers.  The template stores its
        # own automatic number field, which would otherwise duplicate them.
        num_pr = paragraph._p.pPr.find(qn("w:numPr"))
        if num_pr is not None:
            paragraph._p.pPr.remove(num_pr)
        for run in paragraph.runs:
            if run.text:
                set_run_font(run, name=font, size=size, color="000000", bold=bold)

    for table in doc.tables:
        table._tbl.tblPr.clear()
        table._tbl.tblPr.extend(deepcopy(table_template._tbl.tblPr))
        for row in table.rows:
            for cell in row.cells:
                cell._tc.get_or_add_tcPr().clear()
                cell._tc.tcPr.extend(deepcopy(table_template.cell(0, 0)._tc.tcPr))


def ensure_list_style(doc):
    if "List Bullet" not in [style.name for style in doc.styles]:
        style = doc.styles.add_style("List Bullet", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]


def ensure_heading_styles(doc):
    names = {style.name for style in doc.styles}
    for name in ("Heading 2", "Heading 3"):
        if name not in names:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = doc.styles["Normal"]


def build_design():
    target = OUTPUT_DIR / f"{SOFTWARE_NAME}-设计说明书v1.0.docx"
    template = DESIGN_LAYOUT_TEMPLATE
    if not template.exists():
        template = TEMPLATE_DIR / "面向应急救援的无人机异构双网协同规划软件-设计说明书v1.0.docx"
    shutil.copy2(template, target)
    doc = Document(target)
    # The backup is the layout master: retain its cover and front matter, and
    # use its original body paragraphs as direct-format exemplars.
    body_template = doc.paragraphs[17]
    section_template = doc.paragraphs[15]
    subsection_template = doc.paragraphs[16]
    table_template = doc.tables[1]
    set_paragraph_text(doc.paragraphs[0], SOFTWARE_NAME.rsplit(" V1.0", 1)[0], color="000000", size=18, bold=True, font="黑体")
    set_paragraph_text(doc.paragraphs[1], "（V1.0）", color="000000", size=18, bold=True, font="黑体")
    set_paragraph_text(doc.paragraphs[12], "申请单位：[待确认]    [待确认]年[待确认]月[待确认]日", color="000000", size=10)
    replace_template_toc(doc, [
        ("1. 软件介绍", 1, 5), ("1.1 开发目的", 2, 5), ("1.2 面向领域", 2, 5),
        ("1.3 软件的主要功能", 2, 6), ("1.4 软件的技术特点", 2, 6),
        ("2. 软件开发信息", 1, 7), ("3. 开发与运行环境", 1, 7),
        ("3.1 Python环境", 2, 7), ("3.2 C++与仿真环境", 2, 7),
        ("3.3 数据与模型依赖", 2, 8), ("4. 软件架构", 1, 8),
        ("4.1 软件总体架构", 2, 8), ("4.1.1 数据与模型层", 3, 8),
        ("4.1.2 推理决策层", 3, 9), ("4.1.3 仿真交互层", 3, 9),
        ("4.1.4 实验分析层", 3, 9), ("4.2 系统流程", 2, 10),
        ("5. 软件的详细设计", 1, 11), ("5.1 数据与模型构建", 2, 11),
        ("5.2 改进TOPSIS排序", 2, 12), ("5.3 双重滞后判决", 2, 12),
        ("5.4 风险感知增强", 2, 13), ("5.5 算法伪代码", 2, 13),
        ("5.6 基线与消融", 2, 15), ("5.7 接口与模块协同设计", 2, 15),
        ("5.7.1 C++到Python消息", 3, 15), ("5.7.2 Python到C++消息", 3, 15),
        ("5.7.3 Python绑定", 3, 15), ("5.7.4 命令行接口", 3, 15),
        ("6. UI软件设计与运行操作", 1, 16), ("6.1 命令行运行界面", 2, 16),
        ("6.2 运行结果界面", 2, 17), ("7. 运行设计", 1, 18),
        ("8. 系统测试与性能分析", 1, 18), ("8.1 静态检查", 2, 18),
        ("8.2 功能检查", 2, 19), ("8.3 实现边界", 2, 19),
        ("9. 文件与模块索引", 1, 19), ("10. 说明", 1, 20),
    ])
    # The layout backup keeps one sample heading as the first body paragraph;
    # remove that paragraph and all following sample content.  Locate the
    # direct body child so TOC paragraphs inside the content control are not
    # mistaken for the sample body.
    body = doc._element.body
    direct_paragraphs = [child for child in body if child.tag == qn("w:p")]
    marker = next(child for child in direct_paragraphs if "1 软件介绍" in "".join(child.itertext()))
    clear_body_from_child(doc, list(body).index(marker))
    ensure_list_style(doc)
    ensure_heading_styles(doc)
    start_index = len(doc.paragraphs)
    add_markdown(doc, PROJECT / "softcopyright/design_description.md", skip_leading_title=True)
    for paragraph in list(doc.paragraphs[start_index:]):
        if paragraph.text.strip() == "软件设计说明书（V1.0草案）":
            paragraph._element.getparent().remove(paragraph._element)
    # A few layout masters keep the first sample heading inside the front
    # matter container. Remove such stale headings without touching the new
    # Markdown body appended after start_index.
    heading_re = re.compile(r"^\d(?:\.\d+)*\s+\S")
    for paragraph in list(doc.paragraphs[:start_index]):
        if heading_re.match(paragraph.text.strip()):
            paragraph._element.getparent().remove(paragraph._element)
    apply_template_body_format(doc, start_index, body_template, section_template, subsection_template, table_template)
    if len(doc.paragraphs) > start_index:
        doc.paragraphs[start_index].paragraph_format.page_break_before = True
    doc.save(target)


def add_code_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, name="Consolas", size=8, color="000000")


def build_source():
    target = OUTPUT_DIR / f"{SOFTWARE_NAME}-源程序.docx"
    template = TEMPLATE_DIR / "面向应急救援的无人机异构双网协同规划软件-源程序.docx"
    shutil.copy2(template, target)
    doc = Document(target)
    clear_body(doc)
    add_template_title(doc, SOFTWARE_NAME, "核心源程序文档（完整版本，草案）")
    doc.add_paragraph("本源程序文档沿用软著模板排版，按训练、推理、消息接口、仿真、基线、批处理和绘图顺序收录核心源程序。模型权重、训练数据和NS-3外部工作区列入依赖清单。")
    records = source_inventory()
    for index, record in enumerate(records, start=1):
        p = doc.add_paragraph()
        run = p.add_run(f"{index}. {record['file']}（{record['lines']}行）")
        set_run_font(run, name="黑体", size=11, color="000000", bold=True)
        path = PROJECT / record["file"]
        for number, line in enumerate(path.read_text(encoding="utf-8", errors="replace").splitlines(), start=1):
            add_code_paragraph(doc, f"{number:04d} | {line}")
        doc.add_page_break()
    doc.save(target)
    patch_template_headers(target)


def main():
    backup = backup_current()
    remove_old_generated_files()
    build_summary()
    build_application()
    build_design()
    build_source()
    print(f"Backup created: {backup}")
    print(f"Template-based documents generated in: {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
