#!/usr/bin/env python3
"""按实际 PDF 页码重建设计说明书目录。

用法：/usr/bin/python3 fix_toc.py DOCX PDFTXT
PDFTXT 为 `pdftotext -layout` 的输出。目录页与封面共占前 N 页，
正文首页由脚本自动探测（第一个出现“1 软件介绍”且不含“目录”的页）。
"""
import copy
import re
import sys

from docx import Document
from docx.oxml.ns import qn


def norm(s):
    return re.sub(r"\s+", "", s)


def body_headings(doc):
    """返回正文标题列表 [(para_idx, text, level)]，跳过目录区。"""
    out = []
    started = False
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if not t or not p.runs:
            continue
        if not started:
            # 正文从第一个不带制表符的 "1 软件介绍" 开始
            if t == "1 软件介绍" and "\t" not in p.text:
                started = True
            else:
                continue
        r = p.runs[0]
        sz = r.font.size.pt if r.font.size else None
        m = re.match(r"^(\d+(?:\.\d+){0,2})\s", t)
        if sz == 16.0 and m:
            lvl = m.group(1).count(".") + 1
            out.append((i, t, lvl))
        elif m and m.group(1).count(".") == 2 and len(t) < 40:
            # 5.2.x 三级标题字号可能不是 16pt
            out.append((i, t, 3))
    return out


def page_of(pages, title, first_body_page):
    """在正文页范围内找标题首次出现的页码（1-based，PDF 物理页）。"""
    key = norm(title)
    for n in range(first_body_page, len(pages) + 1):
        if key in norm(pages[n - 1]):
            return n
    return None


def toc_rows(doc):
    """返回目录段落列表（含制表符的那些）。"""
    rows = []
    for i, p in enumerate(doc.paragraphs):
        if "\t" in p.text and re.match(r"^\s*\d+(\.\d+){0,2}\s", p.text):
            rows.append(i)
        elif rows and "\t" not in p.text and p.text.strip():
            break
    return rows


def rewrite_row(p, title, page, template):
    """用 template 段落的 run 结构重写目录行。"""
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    tpl_runs = template.runs
    # 结构: [前导空 run...] title \t [空 run...] page [空 run...]
    def clone(src_run, text):
        new = copy.deepcopy(src_run._element)
        for t_el in new.findall(qn("w:t")):
            new.remove(t_el)
        for br in new.findall(qn("w:tab")):
            new.remove(br)
        p._element.append(new)
        from docx.text.run import Run
        run = Run(new, p)
        run.text = text
        return run

    src_title = next((r for r in tpl_runs if r.text.strip() and "\t" not in r.text), tpl_runs[0])
    src_tab = next((r for r in tpl_runs if r.text == "\t"), src_title)
    src_page = next((r for r in reversed(tpl_runs) if r.text.strip().isdigit()), src_title)

    clone(src_title, title)
    tabrun = clone(src_tab, "")
    tabrun._element.append(tabrun._element.makeelement(qn("w:tab"), {}))
    clone(src_page, str(page))


def main():
    docx_path, txt_path = sys.argv[1], sys.argv[2]
    pages = open(txt_path, encoding="utf-8").read().split("\f")
    doc = Document(docx_path)

    heads = body_headings(doc)
    if not heads:
        raise SystemExit("正文标题定位失败")
    first_body_page = None
    for n, pg in enumerate(pages, 1):
        flat = norm(pg)
        if "1软件介绍" in flat and "目录" not in flat[:80] and "......" not in pg:
            first_body_page = n
            break
    if first_body_page is None:
        raise SystemExit("正文首页定位失败")

    plan = []
    for _, title, lvl in heads:
        pg = page_of(pages, title, first_body_page)
        plan.append((title, lvl, pg))

    rows = toc_rows(doc)
    print("目录行数 %d, 正文标题数 %d, 正文首页 %d" % (len(rows), len(plan), first_body_page))
    for title, lvl, pg in plan:
        print("   L%d %-42s -> %s" % (lvl, title, pg))

    if any(pg is None for _, _, pg in plan):
        raise SystemExit("有标题未在 PDF 中定位到，终止")

    # 先深拷贝一份模板段落，避免改写第一行后模板被清空
    template = copy.deepcopy(doc.paragraphs[rows[0]]._element)
    from docx.text.paragraph import Paragraph
    template = Paragraph(template, doc.paragraphs[rows[0]]._parent)
    # 需要的行数可能多于现有行：不足则在最后一行后复制
    while len(rows) < len(plan):
        last = doc.paragraphs[rows[-1]]
        new_el = copy.deepcopy(last._element)
        last._element.addnext(new_el)
        rows = toc_rows(doc)
    # 多余的行删除
    while len(rows) > len(plan):
        p = doc.paragraphs[rows[-1]]
        p._element.getparent().remove(p._element)
        rows = toc_rows(doc)

    for idx, (title, lvl, pg) in zip(rows, plan):
        rewrite_row(doc.paragraphs[idx], title, pg, template)

    doc.save(docx_path)
    print("saved:", docx_path)


if __name__ == "__main__":
    main()
