#!/usr/bin/env python3
"""按两份模板的“6 系统测试与性能分析”体例重写设计说明书第6章。

基线：backup_20260902/...设计说明书v1.0.docx（不覆盖）
输出：同目录下 ...设计说明书v1.0-ch6rewrite.docx
必须用 /usr/bin/python3 运行（python-docx 只装在系统解释器）。
"""
import copy
import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

BASE_DIR = ("/home/suwen/IBN5100/无人机自组网/软著/"
            "LAAVHA软件著作权材料")
SRC = os.path.join(BASE_DIR, "backup_20260902",
                   "无人机遥感异构网络垂直切换智能决策软件 V1.0-设计说明书v1.0.docx")
DST = os.path.join(BASE_DIR,
                   "无人机遥感异构网络垂直切换智能决策软件 V1.0-设计说明书v1.0-ch6rewrite.docx")
EVID = "/home/suwen/reproduce/softcopyright/evidence"
FIG61 = os.path.join(EVID, "fig6_1_runtime_cost.png")
FIG62 = os.path.join(EVID, "fig6_2_stability.png")

BODY_PT = 10.5
HEAD_PT = 16.0
FIG_W = Inches(5.75)


def _style_run(run, size, bold=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    rf.set(qn("w:eastAsia"), "宋体")
    rf.set(qn("w:ascii"), "Times New Roman")
    rf.set(qn("w:hAnsi"), "Times New Roman")


def _blank(p):
    """清空段落的所有 run，保留段落属性。"""
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    return p


def set_heading(p, text, bold):
    _blank(p)
    _style_run(p.add_run(text), HEAD_PT, bold)
    p.paragraph_format.first_line_indent = 0
    p.alignment = None


def set_body(p, text):
    _blank(p)
    _style_run(p.add_run(text), BODY_PT, False)
    pf = p.paragraph_format
    pf.first_line_indent = Pt(21)
    pf.line_spacing = 1.5
    p.alignment = None


def set_caption(p, text):
    _blank(p)
    _style_run(p.add_run(text), BODY_PT, False)
    p.paragraph_format.first_line_indent = 0
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def new_para_after(anchor, doc):
    """在 anchor 段落之后插入一个新空段落并返回。"""
    from docx.text.paragraph import Paragraph
    new_p = copy.deepcopy(anchor._element)
    for child in list(new_p):
        if not child.tag.endswith("}pPr"):
            new_p.remove(child)
    anchor._element.addnext(new_p)
    return Paragraph(new_p, anchor._parent)


def main():
    import ch6_text as T

    for f in (FIG61, FIG62):
        if not os.path.exists(f):
            raise SystemExit("缺少插图: " + f)

    doc = Document(SRC)
    paras = doc.paragraphs

    # 定位第6章
    i_ch6 = i_61 = i_62 = i_idx = None
    for i, p in enumerate(paras):
        t = p.text.strip()
        if t == "6 系统测试与性能分析":
            i_ch6 = i
        elif t == "6.1 计算成本分析":
            i_61 = i
        elif t == "6.2 稳定性与可靠性":
            i_62 = i
        elif t == "文件与模块索引":
            i_idx = i
    if None in (i_ch6, i_61, i_62, i_idx):
        raise SystemExit("章节定位失败: %s" % [i_ch6, i_61, i_62, i_idx])

    # 1) 修正 5.7.3 中残留的图6.3—6.5 编号为 图5.5—5.7（含正文交叉引用）
    ren = {"图6.3": "图5.5", "图6.4": "图5.6", "图6.5": "图5.7"}
    fixed = 0
    for p in paras[:i_ch6]:
        if not any(k in p.text for k in ren):
            continue
        for r in p.runs:
            if any(k in r.text for k in ren):
                new = r.text
                for k, v in ren.items():
                    new = new.replace(k, v)
                r.text = new
                fixed += 1
    # 正文里"图6.3至图6.5"的交叉引用同步
    for p in paras[:i_ch6]:
        for r in p.runs:
            if "图5.5至图5.7" in r.text or "图6.3至" in r.text:
                r.text = r.text.replace("图6.3至图6.5", "图5.5至图5.7")

    # 2) 第8.3节的悬空引用 -> 第6.3节
    for p in paras[:i_ch6]:
        for r in p.runs:
            if "第8.3节" in r.text:
                r.text = r.text.replace("第8.3节", "第6.3节")

    # 3) 删除旧的 6.1/6.2 正文段（i_61+1 .. i_idx-1 之间的非标题段）
    for p in paras[i_61 + 1:i_idx]:
        t = p.text.strip()
        if t in ("6.2 稳定性与可靠性",):
            continue
        p._element.getparent().remove(p._element)

    # 重新取一次段落表
    paras = doc.paragraphs
    i_61 = next(i for i, p in enumerate(paras) if p.text.strip() == "6.1 计算成本分析")
    i_62 = next(i for i, p in enumerate(paras) if p.text.strip() == "6.2 稳定性与可靠性")

    # 4) 写 6.1：正文 + 图6.1 + 图注
    anchor = paras[i_61]
    for txt in T.P61[:3]:
        anchor = new_para_after(anchor, doc)
        set_body(anchor, txt)
    anchor = new_para_after(anchor, doc)
    _blank(anchor)
    anchor.paragraph_format.first_line_indent = 0
    anchor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    anchor.add_run().add_picture(FIG61, width=FIG_W)
    anchor = new_para_after(anchor, doc)
    set_caption(anchor, T.CAP61)
    anchor = new_para_after(anchor, doc)
    set_body(anchor, T.P61[3])

    # 5) 写 6.2：正文 + 图6.2 + 图注
    i_62 = next(i for i, p in enumerate(doc.paragraphs)
                if p.text.strip() == "6.2 稳定性与可靠性")
    anchor = doc.paragraphs[i_62]
    for txt in T.P62[:2]:
        anchor = new_para_after(anchor, doc)
        set_body(anchor, txt)
    anchor = new_para_after(anchor, doc)
    _blank(anchor)
    anchor.paragraph_format.first_line_indent = 0
    anchor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    anchor.add_run().add_picture(FIG62, width=FIG_W)
    anchor = new_para_after(anchor, doc)
    set_caption(anchor, T.CAP62)
    anchor = new_para_after(anchor, doc)
    set_body(anchor, T.P62[2])

    # 6) 追加 6.3 实现边界
    anchor = new_para_after(anchor, doc)
    set_heading(anchor, T.H_63, False)
    for txt in T.P63:
        anchor = new_para_after(anchor, doc)
        set_body(anchor, txt)

    doc.save(DST)
    print("saved:", DST)
    print("图号修正 run 数:", fixed)


if __name__ == "__main__":
    main()
