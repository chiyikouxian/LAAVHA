#!/usr/bin/env python3
"""修复设计说明书：补回 6.3 实现边界与文件模块索引表，改正运行环境表措辞。

6.3 正文与模块表从 /tmp/ch6_before_swap.docx 复制（该副本含完整第6章）。
不改动其余章节。必须用 /usr/bin/python3 运行。
"""
import copy
import os

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

TARGET = ("/home/suwen/IBN5100/无人机自组网/软著/LAAVHA软件著作权材料/"
          "无人机遥感异构网络垂直切换智能决策软件 V1.0-设计说明书v1.0.docx")
SRC = "/tmp/ch6_before_swap.docx"

BODY_PT = 10.5
HEAD_PT = 16.0

RUNTIME_NEW = ("运行平台：Linux Ubuntu 24.04操作系统；\n"
               "运行可视化界面依赖Python 3.10及tkinter，可在Linux桌面环境运行")


def restore_ch63_and_table(doc):
    """从 SRC 复制 6.3 整节（标题+3段）与模块索引表（标题+表格）。"""
    src = Document(SRC)
    sps = src.paragraphs
    i63 = next(k for k, p in enumerate(sps) if p.text.strip() == "6.3 实现边界")
    iidx = next(k for k, p in enumerate(sps)
                if p.text.strip() == "文件与模块索引")
    # 6.3 标题 + 其后到"文件与模块索引"之前的正文段
    block = [sps[k] for k in range(i63, iidx)]
    src_tbl = src.tables[-1]
    if len(src_tbl.columns) != 3 or len(src_tbl.rows) != 8:
        raise SystemExit("模块表形状异常: %dx%d"
                         % (len(src_tbl.rows), len(src_tbl.columns)))

    ps = doc.paragraphs
    # 锚点：文档最后一个非空段（6.2 末段）
    anchor = None
    for p in reversed(ps):
        if p.text.strip():
            anchor = p
            break
    if anchor is None:
        raise SystemExit("找不到插入锚点")
    if any(p.text.strip() == "6.3 实现边界" for p in ps):
        print("6.3 已存在，跳过")
        return 0, False

    added = 0
    cur = anchor._element
    for sp in block:
        new = copy.deepcopy(sp._element)
        cur.addnext(new)
        cur = new
        added += 1
    # 模块索引表标题段
    new_idx = copy.deepcopy(sps[iidx]._element)
    cur.addnext(new_idx)
    cur = new_idx
    added += 1
    # 表格本体
    new_tbl = copy.deepcopy(src_tbl._element)
    cur.addnext(new_tbl)
    # 表后补一个空段，避免表格紧贴节尾
    tail = copy.deepcopy(sps[iidx]._element)
    for child in list(tail):
        if not child.tag.endswith("}pPr"):
            tail.remove(child)
    new_tbl.addnext(tail)
    return added, True


def fix_runtime_table(doc):
    """去掉运行环境表里本软件不存在的"前端软件"措辞。"""
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if "前端软件" not in cell.text:
                    continue
                p = cell.paragraphs[0]
                for extra in cell.paragraphs[1:]:
                    extra._element.getparent().remove(extra._element)
                for r in list(p.runs):
                    r._element.getparent().remove(r._element)
                first = True
                for line in RUNTIME_NEW.split("\n"):
                    r = p.add_run(line if first else "")
                    if not first:
                        r.add_break()
                        r.text = line
                    r.font.size = Pt(BODY_PT)
                    r.font.name = "Times New Roman"
                    rpr = r._element.get_or_add_rPr()
                    rf = rpr.get_or_add_rFonts()
                    rf.set(qn("w:eastAsia"), "宋体")
                    first = False
                return True
    return False


def main():
    if not os.path.exists(SRC):
        raise SystemExit("缺少源副本: " + SRC)
    doc = Document(TARGET)
    n, ok = restore_ch63_and_table(doc)
    print("补回段落 %d 个，模块表: %s" % (n, "已插入" if ok else "未插入"))
    print("运行环境表改写:", "成功" if fix_runtime_table(doc) else "未找到目标单元格")
    doc.save(TARGET)
    print("saved:", TARGET)


if __name__ == "__main__":
    main()
