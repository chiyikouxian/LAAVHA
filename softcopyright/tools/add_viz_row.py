#!/usr/bin/env python3
"""给文件模块索引表补回"运行可视化"一行（从 design_description.md 第332行取值）。"""
import copy

from docx import Document

TARGET = ("/home/suwen/IBN5100/无人机自组网/软著/LAAVHA软件著作权材料/"
          "无人机遥感异构网络垂直切换智能决策软件 V1.0-设计说明书v1.0.docx")
ROW = ("运行可视化",
       "`softcopyright/tools/laavha_viz/`（6个模块）",
       "动画轨迹与时间序列解析、界面布局绘制、交互回放与插图导出")


def main():
    doc = Document(TARGET)
    t = doc.tables[-1]
    if any("运行可视化" in r.cells[0].text for r in t.rows):
        print("该行已存在，跳过")
        return
    last = t.rows[-1]
    new = copy.deepcopy(last._element)
    last._element.addnext(new)
    row = t.rows[-1]
    for cell, text in zip(row.cells, ROW):
        p = cell.paragraphs[0]
        for extra in cell.paragraphs[1:]:
            extra._element.getparent().remove(extra._element)
        runs = p.runs
        if runs:
            runs[0].text = text
            for r in runs[1:]:
                r._element.getparent().remove(r._element)
        else:
            p.add_run(text)
    doc.save(TARGET)
    print("已补行: %s | %s" % (ROW[0], ROW[1]))
    print("表格现为 %dx%d" % (len(t.rows), len(t.columns)))


if __name__ == "__main__":
    main()
