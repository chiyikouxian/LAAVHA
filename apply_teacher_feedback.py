#!/usr/bin/env python3
"""Apply all teacher feedback"""
from docx import Document
from lxml import etree

doc = Document('/home/suwen/reproduce/物联网学报_LAAVHA小论文.docx')
ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# ===== 1. TITLE =====
new_cn = '面向无人机遥感异构网络的注意力LSTM增强型风险感知垂直切换算法'
new_en = 'Attention-LSTM Enhanced Risk-aware Vertical Handoff Algorithm for UAV Remote Sensing Heterogeneous Networks'

for p in doc.paragraphs:
    if '面向无人机遥感监测' in p.text and '自适应垂直切换' in p.text and len(p.text) < 60:
        for r in p.runs[1:]: r.text = ''
        p.runs[0].text = new_cn
        print(f'1. CN title updated')
        break

for p in doc.paragraphs:
    if 'Adaptive vertical handover' in p.text and 'UAV' in p.text:
        for r in p.runs[1:]: r.text = ''
        p.runs[0].text = new_en
        print(f'1. EN title updated')
        break

# ===== 2. COMPRESS ABSTRACT =====
new_cn_abs = (
    '摘  要：针对无人机遥感异构网络中传统垂直切换算法权重固定、'
    '决策滞后及信道波动适应性不足的问题，提出一种注意力LSTM增强型风险感知垂直切换算法（ALERA）。'
    '算法以堆叠LSTM预测网络状态变化，通过注意力机制生成动态权重，融合改进TOPSIS完成候选排序，'
    '并引入自适应滞后（ADH）与风险敏感TOPSIS（RS-TOPSIS）增强机制——前者使切换阈值和确认窗口'
    '随信道波动自适应调节，后者通过置信下界罚分抑制评分波动网络的误选。'
    '基于ns-3与ns3-ai的550次决策级实验表明，ALERA平均切换仅0.24次，'
    '显著优于TOPSIS-Q（4.04次）、Fuzzy-VHO（8.64次）等8种对比算法；'
    '消融实验与合成压力测试分别验证了时序预测、动态权重及增强机制的协同贡献。'
)
new_en_abs = (
    'Abstract: To address the challenges of fixed weights, decision latency, and poor channel-fluctuation '
    'adaptability in traditional vertical handover algorithms for UAV remote sensing heterogeneous networks, '
    'an Attention-LSTM Enhanced Risk-aware Vertical Handoff Algorithm (ALERA) is proposed. '
    'A stacked LSTM predicts network state changes; an attention mechanism generates dynamic attribute weights; '
    'an improved TOPSIS completes candidate ranking; and two enhanced mechanisms\u2014adaptive hysteresis (ADH) '
    'and risk-sensitive TOPSIS (RS-TOPSIS)\u2014are introduced. ADH auto-adjusts the handover threshold and '
    'confirmation window based on channel volatility, while RS-TOPSIS penalizes networks with volatile scores '
    'via a lower-confidence-bound penalty. Decision-level experiments (550 runs) on ns-3 and ns3-ai show that '
    'ALERA achieves an average of only 0.24 handovers, significantly outperforming eight baselines including '
    'TOPSIS-Q (4.04) and Fuzzy-VHO (8.64). Ablation and synthetic stress tests validate the synergistic '
    'contribution of temporal prediction, dynamic weighting, and the enhanced mechanisms.'
)

for p in doc.paragraphs:
    if p.text.startswith('摘  要：') and '无人机' in p.text:
        for r in p.runs[1:]: r.text = ''
        p.runs[0].text = new_cn_abs
        print('2. CN abstract compressed')
        break
for p in doc.paragraphs:
    if p.text.startswith('Abstract:') and 'UAV' in p.text[:60]:
        for r in p.runs[1:]: r.text = ''
        p.runs[0].text = new_en_abs
        print('2. EN abstract updated')
        break

# ===== 3. KEYWORDS =====
for p in doc.paragraphs:
    if '关键词' in p.text[:10]:
        for r in p.runs[1:]: r.text = ''
        p.runs[0].text = '关键词：无人机遥感；异构网络；垂直切换；LSTM；注意力机制；自适应滞后；风险敏感TOPSIS'
        break
for p in doc.paragraphs:
    if 'Key words' in p.text[:20]:
        for r in p.runs[1:]: r.text = ''
        p.runs[0].text = 'Key words: UAV remote sensing, heterogeneous network, vertical handover, LSTM, attention mechanism, adaptive hysteresis, risk-sensitive TOPSIS'
        break
print('3. Keywords updated')

# ===== 4. SECTION 2 HEADING =====
for p in doc.paragraphs:
    if '2 LAAVHA自适应垂直切换算法' in p.text:
        for r in p.runs:
            if 'LAAVHA' in r.text:
                r.text = r.text.replace('LAAVHA', 'ALERA')
                print('4. Section 2 heading updated')
                break

# ===== 5. INTRODUCTION RESTRUCTURE =====
# Find intro paragraphs
intro_paras = []
in_intro = False
for i, p in enumerate(doc.paragraphs):
    if '0 引言' in p.text:
        in_intro = True
        continue
    if in_intro and '1 系统模型' in p.text:
        break
    if in_intro and len(p.text) > 50:
        intro_paras.append(i)

print(f'5. Found {len(intro_paras)} intro body paragraphs at indices: {intro_paras}')

# PARA 1 (merged): Traditional methods in UAV remote sensing - limitations
# Original P[14] (2677 chars) + P[15] (?) 
# We merge P1+P2 into one paragraph about limitations of traditional methods
p1_idx = intro_paras[0]
p1_new = (
    '无人机遥感系统在灾害评估、农业普查、环境监测等领域广泛应用，无人机搭载的高分辨率光学/多光谱载荷'
    '在巡航过程中产生大量遥感图像数据，需通过5G、LTE和WiFi等异构无线网络实时回传至地面站[1-3]。'
    '然而，无人机的高机动性导致其频繁穿越不同网络的覆盖边界，触发垂直切换[4]。'
    '传统垂直切换研究主要围绕多属性决策方法展开，如TOPSIS[7-8]、VIKOR[22]、GRA[22]、COPRAS[22]、'
    'SPOTIS[22]、Fuzzy-VHO[4]和SAW等，通过构建信噪比、时延、吞吐量等网络属性的评价矩阵进行候选排序。'
    '这些方法存在两个共性问题：一是权重通常由熵权法或层次分析法一次性确定[10-11]，无法随飞行阶段'
    '和信道环境动态调整；二是仅依据当前时刻的网络指标做决策，缺乏对未来状态变化的前瞻性判断，'
    '当无人机高速飞向WiFi覆盖边缘时，可能在链路实际中断后才触发切换[12-13]。'
    '此外，近年深度强化学习（DRL）方法被引入切换决策[12-14]，但其黑箱性质和训练不稳定性在安全攸关的'
    '遥感巡航场景中带来可解释性和可靠性方面的顾虑。'
)
for r in doc.paragraphs[p1_idx].runs[1:]: r.text = ''
doc.paragraphs[p1_idx].runs[0].text = p1_new
print(f'5. P{intro_paras[0]}: P1+P2 merged (traditional methods limitations)')

# PARA 2 (new): LSTM applications and limitations
lstm_idx = intro_paras[1]
lstm_new = (
    '长短期记忆网络（LSTM）通过门控结构有效缓解梯度消失问题[15]，具有捕获长时间序列依赖关系的能力，'
    '已被应用于5G异构网络的切换判决和网络状态预测[16-17]。然而，单一LSTM网络存在两方面局限：'
    '其一，LSTM将候选网络的5维状态指标视为独立的时间序列分别建模，无法显式捕获指标之间的耦合关系'
    '——例如吞吐量与丢包率的负相关性、SINR与时延的关联性等；'
    '其二，LSTM输出的预测状态仍需配合固定或人工设定的属性权重才能完成多属性决策，无法根据飞行阶段'
    '和移动状态自适应调整各指标的重要性。'
)
for r in doc.paragraphs[lstm_idx].runs[1:]: r.text = ''
doc.paragraphs[lstm_idx].runs[0].text = lstm_new
print(f'5. P{intro_paras[1]}: LSTM applications + limitations (formerly P3)')

# PARA 3 (new): Two-level summary
p3_idx = intro_paras[2]
p3_new = (
    '综上所述，当前无人机遥感异构网络中的垂直切换面临两个层面的挑战：'
    '在算法层面，传统MADM方法权重固定、缺乏时序预测能力，单一LSTM即使引入预测机制仍无法自适应'
    '调节属性权重以适应飞行阶段变化；在场景层面，无人机巡航过程中信道环境动态变化——平稳飞行时'
    '信号波动小，跨越覆盖边界或遭遇遮挡时信号剧烈振荡，固定的切换阈值和确认窗口难以在两者之间取得平衡。'
    '针对第一个层面，引入注意力机制对候选网络状态矩阵进行自注意力建模，使模型自主学习属性间交互关系'
    '并输出随飞行阶段变化的动态权重；针对第二个层面，在决策层引入自适应滞后（ADH）与风险敏感TOPSIS'
    '（RS-TOPSIS）增强机制，使切换参数根据信道波动和飞行状态自适应调节。'
)
for r in doc.paragraphs[p3_idx].runs[1:]: r.text = ''
doc.paragraphs[p3_idx].runs[0].text = p3_new
print(f'5. P{intro_paras[2]}: Two-level limitation summary')

# PARA 4 (new): Our proposal
p4_idx = intro_paras[3]
p4_new = (
    '针对以上问题，本文提出注意力LSTM增强型风险感知垂直切换算法（ALERA），'
    '在LSTM预测与TOPSIS决策框架之上引入注意力机制与两阶段增强策略。'
    '算法采用"预测-权重-决策"三阶段架构：堆叠LSTM预测候选网络短期状态变化，'
    '注意力机制以当前状态矩阵为自注意力输入融合移动特征生成动态权重，'
    '融合当前与预测状态构建改进TOPSIS决策矩阵完成候选排序。'
    '在此基础上，自适应滞后（ADH）机制将固定切换阈值Δ_th和确认窗口T替换为上下文感知的可变参数'
    '——以服务网络近5个周期的SINR标准差度量信道波动，阈值随波动从0.03线性增长至0.08，'
    '确认窗口随飞行速度在[2,4]范围内自适应缩短；风险敏感TOPSIS（RS-TOPSIS）机制引入置信下界罚分，'
    '以C_i^robust = C_i − 0.5·σ_i^C作为排序依据，对评分波动较大的候选网络施加惩罚，'
    '使排序偏好倾向历史表现更稳定的网络。两项增强机制均作用于决策层，不改变模型结构。'
)
for r in doc.paragraphs[p4_idx].runs[1:]: r.text = ''
doc.paragraphs[p4_idx].runs[0].text = p4_new
print(f'5. P{intro_paras[3]}: Proposal with innovations')

# Delete any remaining old intro paragraphs beyond P4
if len(intro_paras) > 4:
    for idx in intro_paras[4:]:
        p = doc.paragraphs[idx]
        p._element.getparent().remove(p._element)
        print(f'5. Deleted extra intro P{idx}')

# ===== 6. MOVE sentence from intro to 3.5 =====
# Find "何时切换/切向谁" in any remaining intro paragraph and remove
for p in doc.paragraphs:
    if '两项机制分别决定' in p.text and '何时切换' in p.text and '切向谁' in p.text:
        for r in p.runs:
            if '两项机制分别决定' in r.text:
                # Remove just that sentence, keep the rest
                old = '两项机制分别决定"何时切换"与"切向谁"，在不改变模型结构的前提下为算法赋予信道波动感知能力。'
                r.text = r.text.replace(old, '')
                print('6. Removed sentence from intro')
                break

# Add to 3.5 section end
for p in doc.paragraphs:
    if '该实验清晰证明：两项增强机制在不牺牲平稳场景响应速度的前提下' in p.text:
        for r in p.runs:
            if '有效抑制瞬时异常触发的乒乓切换。' in r.text:
                r.text = r.text.replace(
                    '有效抑制瞬时异常触发的乒乓切换。',
                    '有效抑制瞬时异常触发的乒乓切换。两项机制分别决定"何时切换"与"切向谁"，在不改变模型结构的前提下为算法赋予信道波动感知能力。'
                )
                print('6. Added sentence to 3.5 section end')
                break

doc.save('/home/suwen/reproduce/物联网学报_LAAVHA小论文.docx')
print('\nAll changes saved.')
