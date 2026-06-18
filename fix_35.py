from docx import Document
doc = Document('/home/suwen/reproduce/物联网学报_LAAVHA小论文_修订版.docx')

# Find 3.5
idx_35 = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if '3.5 增强机制验证' in t:
        idx_35 = i
        break

if idx_35 is None:
    print("3.5 not found")
    for i, p in enumerate(doc.paragraphs):
        if '增强机制验证' in p.text:
            print(f"P[{i}]: [{p.text[:80]}]")
    exit()

p_content = doc.paragraphs[idx_35 + 1]
new_text = (
    "为验证两项增强机制的有效性，设计了两组靶向实验。"
    "实验一（自适应滞后验证）：在seed=100的原始仿真数据基础上，对t=3~6s区间注入5G SINR ±20dB正弦振荡、"
    "同步劣化时延/吞吐量/丢包率，并将WiFi信号提升10dB，构造信道剧烈波动场景。"
    "图5以4行子图对比了自适应滞后机制与固定低阈值（Δ_th=0.03）的表现："
    "第1行展示SINR振荡过程；第2行展示固定低阈值方案在振荡期间触发了误切换（红色虚线）；"
    "第3行展示LAAVHA-enhanced的自适应阈值在检测到SINR波动度上升后自动抬升至约0.07，"
    "成功抑制了误切换；第4行展示自适应阈值随时间的动态变化——"
    "波动区间内升高、波动消退后恢复至基础值，验证了自适应参数的可逆性。"
    "该实验表明，自适应滞后机制能够在信道剧烈波动时自动提升切换门槛，抑制乒乓切换。"
    "实验二（风险敏感TOPSIS验证）：构造WiFi（均值0.55，标准差0.04）与LTE（均值0.52，标准差0.015）"
    "的100步贴近度竞争场景，对比原始TOPSIS与LCB-TOPSIS的排序差异。"
    "结果表明原始TOPSIS在WiFi峰值期频繁选择WiFi（83/100次），"
    "而LCB-TOPSIS通过σ罚分将WiFi选择降至72次、LTE从15次升至27次——"
    "体现了'优先选择评分稳定网络'的排序偏好，与遥感数据传输对连续性的需求一致。"
    "两项实验共同验证了增强机制的设计目标：在不重新训练模型的条件下，"
    "仅通过决策层参数和排序逻辑的调整即可为算法赋予信道波动感知能力。"
    "需要指出的是，在当前代理5G场景下候选网络贴近度差距较大，"
    "自适应滞后与风险敏感机制的完整优势需在真实NR协议栈引入评分竞争区间后才能充分释放，"
    "当前的实验主要提供了机制级的功能验证。"
)
for r in p_content.runs[1:]:
    r.text = ''
p_content.runs[0].text = new_text
print("Content updated")

# Update figure 5 caption
for p in doc.paragraphs:
    if '图5' in p.text and ('LAAVHA与LAAVHA-enhanced' in p.text or '振荡' in p.text):
        for r in p.runs:
            if '图5' in r.text[:10]:
                r.text = '图5 自适应滞后机制验证：固定低阈值(0.03)与LAAVHA-enhanced在±20dB SINR振荡场景下的对比'
                print("Caption updated")
                break
        break

# Also update any reference to 图5第四个子图 since figure layout changed
for p in doc.paragraphs:
    if '图5第四个子图' in p.text:
        for r in p.runs:
            if '图5第四个子图' in r.text:
                r.text = r.text.replace('图5第四个子图', '图5第4行')
                print("Fixed 图5 reference")

doc.save('/home/suwen/reproduce/物联网学报_LAAVHA小论文_修订版.docx')
print("Done")
